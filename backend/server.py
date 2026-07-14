"""NivXRay — FastAPI backend."""
from __future__ import annotations
import os
import re
import json
import base64
import logging
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

import bcrypt
import jwt
import asyncio
from dotenv import load_dotenv
from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, EmailStr

from operations import (
    OPERATIONS, list_operations, run_operation,
    extract_iocs, detect_payload_type, mitre_map, yara_lite_scan, risk_score,
)
import ops_extended  # noqa: F401  — registers +40 new operations into OPERATIONS
from smart_decoder import smart_decode
from magic_decoder import magic_decode, score_output as magic_score
from osint import enrich_iocs, OSINT_SERVICES
from feeds import SOURCES as FEED_SOURCES, sync_source
from lolbas import (
    scan_lolbas, load_from_db as lolbas_load, maybe_refresh as lolbas_maybe_refresh,
    refresh_from_source as lolbas_refresh, get_status as lolbas_status,
)
import models_studio as ms
import sample_library as sl

from emergentintegrations.llm.chat import LlmChat, UserMessage

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

MONGO_URL = os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
JWT_SECRET = os.environ["JWT_SECRET"]
ADMIN_EMAIL = os.environ["ADMIN_EMAIL"]
ADMIN_PASSWORD = os.environ["ADMIN_PASSWORD"]
EMERGENT_LLM_KEY = os.environ["EMERGENT_LLM_KEY"]

JWT_ALG = "HS256"
JWT_EXPIRE_HOURS = 24 * 7

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

app = FastAPI(title="NivXRay API")
api = APIRouter(prefix="/api")
security = HTTPBearer()

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
log = logging.getLogger("nivxray")


# =============================================================================
# Models
# =============================================================================
class LoginIn(BaseModel):
    email: EmailStr
    password: str


class TokenOut(BaseModel):
    access_token: str
    token_type: str = "bearer"
    email: str


class RecipeStep(BaseModel):
    op: str
    args: Dict[str, Any] = Field(default_factory=dict)


class RunRecipeIn(BaseModel):
    input: str
    steps: List[RecipeStep] = Field(default_factory=list)


class RunRecipeOut(BaseModel):
    output: str
    steps_output: List[Dict[str, Any]] = Field(default_factory=list)
    detected_type: Optional[Dict[str, str]] = None
    errors: List[Dict[str, str]] = Field(default_factory=list)


class AutoIn(BaseModel):
    input: str


class AnalyzeIn(BaseModel):
    input: str
    output: Optional[str] = None
    use_ai_verdict: bool = False
    enrich_osint: bool = True
    describe: bool = False
    persona_id: Optional[str] = None
    provider_id: Optional[str] = None


class TroubleshootIn(BaseModel):
    input: str
    steps: List[RecipeStep] = Field(default_factory=list)
    error: Optional[str] = None


class ShareIn(BaseModel):
    input: str
    steps: List[RecipeStep] = Field(default_factory=list)


class SettingsUpdateIn(BaseModel):
    keys: Dict[str, str] = Field(default_factory=dict)  # {service_id: api_key}


# =============================================================================
# Auth
# =============================================================================
def _hash_password(pw: str) -> str:
    return bcrypt.hashpw(pw.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def _verify_password(pw: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


def _create_token(email: str) -> str:
    payload = {
        "sub": email,
        "iat": datetime.now(timezone.utc),
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_HOURS),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)


async def get_current_user(creds: HTTPAuthorizationCredentials = Depends(security)) -> Dict[str, Any]:
    try:
        payload = jwt.decode(creds.credentials, JWT_SECRET, algorithms=[JWT_ALG])
        email = payload.get("sub")
        if not email:
            raise HTTPException(status_code=401, detail="Invalid token")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    user = await db.users.find_one({"email": email}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


async def require_admin(user=Depends(get_current_user)) -> Dict[str, Any]:
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


async def seed_admin():
    existing = await db.users.find_one({"email": ADMIN_EMAIL})
    if existing:
        return
    await db.users.insert_one({
        "email": ADMIN_EMAIL,
        "password": _hash_password(ADMIN_PASSWORD),
        "role": "admin",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    log.info(f"Seeded admin user: {ADMIN_EMAIL}")


# =============================================================================
# Settings
# =============================================================================
async def load_osint_keys() -> Dict[str, str]:
    doc = await db.settings.find_one({"_id": "osint_keys"}) or {}
    return doc.get("keys", {})


def _mask(v: str) -> str:
    if not v: return ""
    if len(v) <= 8: return "•" * len(v)
    return v[:4] + "•" * max(4, len(v) - 8) + v[-4:]


# =============================================================================
# LLM helpers
# =============================================================================
def _new_chat(session_id: str, system: str,
              provider: str = "anthropic", model: str = "claude-sonnet-4-5-20250929") -> LlmChat:
    return LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=session_id,
        system_message=system,
    ).with_model(provider, model)


async def _llm_json(session_id: str, system: str, user: str, retries: int = 1,
                    provider: str = "anthropic", model: str = "claude-sonnet-4-5-20250929") -> Dict[str, Any]:
    chat = _new_chat(session_id, system, provider=provider, model=model)
    last_err = None
    for _ in range(retries + 1):
        try:
            resp = await chat.send_message(UserMessage(text=user))
            raw = resp if isinstance(resp, str) else str(resp)
            cleaned = raw.strip()
            m = re.search(r"```(?:json)?\s*(.*?)\s*```", cleaned, re.DOTALL)
            if m: cleaned = m.group(1)
            if not cleaned.lstrip().startswith("{"):
                m2 = re.search(r"\{.*\}", cleaned, re.DOTALL)
                if m2: cleaned = m2.group(0)
            return json.loads(cleaned)
        except Exception as e:
            last_err = e
    raise HTTPException(status_code=502, detail=f"LLM JSON parse failed: {last_err}")


async def _llm_text(session_id: str, system: str, user: str) -> str:
    chat = _new_chat(session_id, system)
    try:
        r = await chat.send_message(UserMessage(text=user))
        return r if isinstance(r, str) else str(r)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"LLM error: {e}")


# =============================================================================
# Endpoints — Auth
# =============================================================================
@api.get("/")
async def root():
    return {"service": "NivXRay", "status": "ok"}


@api.post("/auth/login", response_model=TokenOut)
async def login(body: LoginIn):
    u = await db.users.find_one({"email": body.email})
    if not u or not _verify_password(body.password, u["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return TokenOut(access_token=_create_token(body.email), email=body.email)


@api.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return user


# =============================================================================
# Endpoints — Operations & Recipes
# =============================================================================
@api.get("/operations")
async def get_ops(user=Depends(get_current_user)):
    return list_operations()


@api.get("/examples")
async def get_examples(user=Depends(get_current_user)):
    return EXAMPLES


@api.post("/recipe/run", response_model=RunRecipeOut)
async def run_recipe(body: RunRecipeIn, user=Depends(get_current_user)):
    current = body.input
    steps_output: List[Dict[str, Any]] = []
    errors: List[Dict[str, str]] = []
    for i, step in enumerate(body.steps):
        try:
            current = run_operation(step.op, current, step.args)
            steps_output.append({
                "index": i, "op": step.op,
                "output_preview": current[:400],
                "output_length": len(current),
            })
        except Exception as e:
            errors.append({"index": str(i), "op": step.op, "error": str(e)})
            steps_output.append({"index": i, "op": step.op, "error": str(e)})
            break
    return RunRecipeOut(
        output=current, steps_output=steps_output,
        detected_type=detect_payload_type(current), errors=errors,
    )


@api.post("/upload")
async def upload(file: UploadFile = File(...), user=Depends(get_current_user)):
    """Universal file upload — accepts ANY file format.

    Returns:
      - filename, size
      - hashes (MD5 / SHA1 / SHA256)
      - detected file type (via magic bytes)
      - text (if the file is textual)
      - hex_dump (first 512 bytes)
      - strings (printable strings >=4 chars, first 200)
      - base64 (full content base64, for very small binaries)
      - content (best-effort text for the Input box)
    """
    import hashlib
    raw = await file.read()
    size = len(raw)

    hashes = {
        "md5": hashlib.md5(raw).hexdigest(),
        "sha1": hashlib.sha1(raw).hexdigest(),
        "sha256": hashlib.sha256(raw).hexdigest(),
    }

    # detect file type via magic bytes
    file_type = _detect_file_type(raw, file.filename or "")

    # try to decode as text
    text = None
    try:
        candidate = raw.decode("utf-8")
        if _mostly_printable(candidate):
            text = candidate
    except UnicodeDecodeError:
        pass
    if text is None:
        try:
            candidate = raw.decode("utf-16-le")
            if _mostly_printable(candidate):
                text = candidate
        except UnicodeDecodeError:
            pass

    # hex dump (first 512 bytes)
    hex_dump = _hex_dump(raw[:512])

    # extract printable strings (>= 4 chars)
    strings_out = _extract_strings(raw, min_len=4, limit=400)

    # what to put in the Input box
    if text is not None:
        content = text[:400_000]
    else:
        content = (
            f"[BINARY FILE — {file.filename}]\n"
            f"Size: {size} bytes\n"
            f"Type: {file_type['label']}\n"
            f"MD5:    {hashes['md5']}\n"
            f"SHA1:   {hashes['sha1']}\n"
            f"SHA256: {hashes['sha256']}\n\n"
            f"── HEX DUMP (first 512 bytes) ──\n{hex_dump}\n\n"
            f"── EXTRACTED STRINGS (top {min(200, len(strings_out))}) ──\n"
            + "\n".join(strings_out[:200])
        )

    return {
        "filename": file.filename,
        "size": size,
        "hashes": hashes,
        "file_type": file_type,
        "text": text[:400_000] if text else None,
        "hex_dump": hex_dump,
        "strings": strings_out,
        "content": content,
    }


def _detect_file_type(raw: bytes, filename: str) -> Dict[str, str]:
    """Detect file type from magic bytes + extension fallback."""
    magics = [
        (b"MZ", "PE (Windows executable / DLL)", "application/x-dosexec"),
        (b"\x7fELF", "ELF (Linux executable)", "application/x-elf"),
        (b"\xCA\xFE\xBA\xBE", "Java class / Mach-O fat", "application/java-vm"),
        (b"\xFE\xED\xFA", "Mach-O binary", "application/x-mach-binary"),
        (b"PK\x03\x04", "ZIP archive (docx/xlsx/jar/apk possible)", "application/zip"),
        (b"Rar!\x1a\x07", "RAR archive", "application/vnd.rar"),
        (b"\x1f\x8b", "GZIP compressed", "application/gzip"),
        (b"\x42\x5a\x68", "BZIP2 compressed", "application/x-bzip2"),
        (b"\xFD7zXZ", "XZ compressed", "application/x-xz"),
        (b"%PDF-", "PDF document", "application/pdf"),
        (b"\xD0\xCF\x11\xE0", "MS OLE compound (legacy Office / MSI)", "application/x-ole"),
        (b"\x89PNG", "PNG image", "image/png"),
        (b"\xff\xd8\xff", "JPEG image", "image/jpeg"),
        (b"GIF87a", "GIF image", "image/gif"),
        (b"GIF89a", "GIF image", "image/gif"),
        (b"#!/", "Shell script (shebang)", "text/x-shellscript"),
        (b"<?xml", "XML document", "application/xml"),
        (b"{\"", "JSON (likely)", "application/json"),
    ]
    for prefix, label, mime in magics:
        if raw.startswith(prefix):
            return {"label": label, "mime": mime, "extension": _ext(filename)}
    # heuristics
    if _mostly_printable(raw[:2048].decode("utf-8", errors="replace")):
        return {"label": "Plain text", "mime": "text/plain", "extension": _ext(filename)}
    return {"label": "Unknown binary", "mime": "application/octet-stream", "extension": _ext(filename)}


def _ext(filename: str) -> str:
    if "." not in filename: return ""
    return filename.rsplit(".", 1)[-1].lower()


def _mostly_printable(s: str, threshold: float = 0.85) -> bool:
    if not s: return False
    printable = sum(1 for c in s if c.isprintable() or c in "\n\r\t")
    return printable / max(1, len(s)) >= threshold


def _hex_dump(data: bytes, width: int = 16) -> str:
    lines = []
    for i in range(0, len(data), width):
        chunk = data[i:i + width]
        hex_part = " ".join(f"{b:02x}" for b in chunk)
        ascii_part = "".join(chr(b) if 32 <= b < 127 else "." for b in chunk)
        lines.append(f"{i:08x}  {hex_part:<{width * 3}}  {ascii_part}")
    return "\n".join(lines)


def _extract_strings(raw: bytes, min_len: int = 4, limit: int = 400) -> List[str]:
    out = []
    cur = []
    for b in raw:
        if 32 <= b < 127:
            cur.append(chr(b))
        else:
            if len(cur) >= min_len:
                out.append("".join(cur))
                if len(out) >= limit: break
            cur = []
    if len(cur) >= min_len and len(out) < limit:
        out.append("".join(cur))
    return out


# =============================================================================
# Endpoints — Smart deterministic Auto-Decode (no AI needed)
# =============================================================================
class MagicIn(BaseModel):
    input: str
    max_depth: int = 4
    max_branches: int = 3
    top_n: int = 3


@api.post("/decode/magic")
async def decode_magic(body: MagicIn, user=Depends(get_current_user)):
    """Recursive multi-branch auto-decoder — returns top-N candidate chains + scores."""
    return magic_decode(body.input, max_depth=body.max_depth,
                        max_branches=body.max_branches, top_n=body.top_n)


class ShellcodeIn(BaseModel):
    input: str = Field(..., description="Hex, base64, or raw bytes to analyse")
    arch: Optional[str] = Field(None, description="Arch hint: x86 / x86_64 / arm / arm64 / thumb")
    max_insns: int = 300


class CommandAnalyzeIn(BaseModel):
    input: str = Field(..., description="Raw command line to semantically analyse")
    force_decode_span: Optional[str] = Field(
        None, description="If the previous /analyze/command call returned "
                          "needs_choice=true, resubmit with the chosen span here."
    )


@api.post("/analyze/command")
async def analyze_command_endpoint(body: CommandAnalyzeIn, user=Depends(get_current_user)):
    """Intelligent Command-Line Analysis Engine — semantic parsing first, then
    decode ONLY the identified high-confidence payload span(s).

    Returns a full analyst report: parsed structure, identified payloads with
    confidence scores, decode chains, extracted IOCs, LOLBin detection, MITRE
    ATT&CK mapping and a behavior summary. If two candidates tie within 0.05
    confidence, `needs_choice:true` is set — resubmit with `force_decode_span`.
    """
    from command_analyzer import analyze_command as _ac
    return _ac(body.input, force_decode_span=body.force_decode_span)


@api.post("/analyze/shellcode")
async def analyze_shellcode(body: ShellcodeIn, user=Depends(get_current_user)):
    """Shellcode / binary analysis — auto-detects architecture, disassembles via
    Capstone, extracts IOCs and API imports. Accepts hex, base64, or raw utf-8.

    This is the "stop condition" endpoint for the recursive decode-and-route
    pipeline: when text-decoding produces a high-entropy binary buffer, hand it
    here instead of continuing to decode.
    """
    from shellcode_analyzer import analyze as _analyze_shellcode
    import base64 as _b64

    raw_in = body.input.strip()
    data: bytes = b""
    src = "utf8"
    # Try hex first (even-length, all hex chars, whitespace tolerated)
    hex_stripped = re.sub(r"[\s:]", "", raw_in)
    if hex_stripped and len(hex_stripped) % 2 == 0 and re.fullmatch(r"[0-9a-fA-F]+", hex_stripped):
        try:
            data = bytes.fromhex(hex_stripped)
            src = "hex"
        except Exception:
            data = b""
    # Fall back to base64
    if not data:
        try:
            b64 = re.sub(r"\s+", "", raw_in)
            data = _b64.b64decode(b64 + "=" * (-len(b64) % 4), validate=False)
            if data:
                src = "base64"
        except Exception:
            data = b""
    if not data:
        data = raw_in.encode("utf-8", errors="replace")
        src = "utf8"

    result = _analyze_shellcode(data, arch=body.arch, max_insns=body.max_insns)
    result["input_source"] = src
    result["input_bytes"] = len(data)
    return result



@api.post("/decode/smart")
async def decode_smart(body: AutoIn, user=Depends(get_current_user)):
    # 1. Check custom decode recipes first — if any regex matches, apply that
    #    recipe on top of the default deterministic pipeline as an ADVISORY hint.
    custom_matches = await ms.find_matching_recipes(db, body.input)
    if custom_matches:
        # apply first (highest-usage) matching recipe; record usage
        best = custom_matches[0]
        await ms.increment_usage(db, best["id"])
        # execute the custom recipe
        try:
            steps_for_run = [RecipeStep(op=s["op"], args=s.get("args") or {}) for s in best["ops"] if s.get("op") in OPERATIONS]
            if steps_for_run:
                result_custom = await run_recipe(
                    RunRecipeIn(input=body.input, steps=steps_for_run), user=user,
                )
                recipe_out = [
                    {"op": s.op, "args": s.args, "reason": f"custom recipe: {best['name']}",
                     "custom": True, "model_id": best["id"], "model_name": best["name"]}
                    for s in steps_for_run
                ]
                return {
                    "recipe": recipe_out,
                    "output": result_custom.output,
                    "notes": [f"Applied custom recipe '{best['name']}' from Model Studio"],
                    "detected_type": detect_payload_type(result_custom.output),
                    "custom_recipes_matched": [
                        {"id": r["id"], "name": r["name"]} for r in custom_matches
                    ],
                }
        except Exception as e:
            log.warning("custom recipe '%s' failed, falling back to smart_decode: %s", best.get("name"), e)

    result = smart_decode(body.input)
    return {
        "recipe": [{"op": s["op"], "args": s.get("args", {}), "reason": s["reason"]} for s in result["steps"]],
        "output": result["output"],
        "notes": result["notes"],
        "detected_type": detect_payload_type(result["output"]),
        "custom_recipes_matched": [
            {"id": r["id"], "name": r["name"]} for r in custom_matches
        ],
    }


# =============================================================================
# Endpoints — Threat Analysis
# =============================================================================
@api.post("/analyze")
async def analyze(body: AnalyzeIn, user=Depends(get_current_user)):
    text = (body.output or "") + "\n" + body.input
    iocs = extract_iocs(text)
    mitre = mitre_map(text)
    yara = yara_lite_scan(text)
    lolbas = scan_lolbas(text)
    risk = risk_score(mitre, yara, iocs)

    # cross-reference against local threat-intel database
    ti_hits = await _lookup_ti_hits(iocs)

    osint_data = None
    if body.enrich_osint:
        try:
            keys = await load_osint_keys()
            osint_data = await enrich_iocs(iocs, keys)
        except Exception as e:
            osint_data = {"error": str(e)}

    ai_verdict = None
    description = None
    if body.use_ai_verdict or body.describe:
        try:
            ai_bundle = await _ai_describe_and_verdict(
                body.input, body.output or "", iocs, mitre, yara, osint_data or {},
                lolbas=lolbas,
                want_verdict=body.use_ai_verdict, want_describe=body.describe,
            )
            ai_verdict = ai_bundle.get("verdict") if body.use_ai_verdict else None
            description = ai_bundle.get("description") if body.describe else None
        except Exception as e:
            if body.use_ai_verdict: ai_verdict = {"error": str(e)}
            if body.describe: description = {"error": str(e)}

    # Merge AI-derived MITRE techniques with heuristic ones (AI takes priority; dedupe by ID)
    merged_mitre = list(mitre)
    if description and isinstance(description, dict):
        ai_mitre = description.get("mitre_techniques") or []
        seen_ids = {m["id"] for m in merged_mitre}
        for m in ai_mitre:
            if isinstance(m, dict) and m.get("id") and m["id"] not in seen_ids:
                merged_mitre.append({
                    "id": m["id"],
                    "technique": m.get("technique", ""),
                    "tactic": m.get("tactic", ""),
                    "evidence": m.get("evidence", ""),
                    "source": "ai",
                })
                seen_ids.add(m["id"])
        # tag heuristic ones
        for m in merged_mitre:
            m.setdefault("source", "heuristic")

    return {
        "iocs": iocs, "mitre": merged_mitre, "yara": yara, "lolbas": lolbas, "risk": risk,
        "osint": osint_data, "ti_hits": ti_hits,
        "ai_verdict": ai_verdict, "description": description,
    }


async def _lookup_ti_hits(iocs: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    """Cross-reference extracted IOCs against local Threat-Intel DB."""
    values: List[str] = []
    for k in ("urls", "ips", "domains", "md5", "sha1", "sha256"):
        values.extend(iocs.get(k) or [])
    if not values:
        return []
    hits = []
    async for doc in db.iocs.find({"value": {"$in": values}}, {"_id": 0}):
        hits.append(doc)
    return hits


# =============================================================================
# SSE streaming — /api/analyze/stream
# Streams the analysis pipeline live so long-running AI investigations don't
# idle out on Cloudflare (60s). Emits named events for each phase and a final
# `result` event containing the complete analysis payload.
# =============================================================================
def _sse(event: str, payload: Any) -> bytes:
    body = json.dumps(payload, default=str)
    return f"event: {event}\ndata: {body}\n\n".encode("utf-8")


@api.post("/analyze/stream")
async def analyze_stream(body: AnalyzeIn, user=Depends(get_current_user)):
    """SSE analog of /api/analyze — streams progress + partial + final result."""

    async def gen():
        # 1) Deterministic extraction (fast)
        try:
            yield _sse("status", {"phase": "extract", "message": "Extracting IOCs, MITRE, YARA, LOLBAS…"})
            text = (body.output or "") + "\n" + body.input
            iocs = extract_iocs(text)
            mitre = mitre_map(text)
            yara = yara_lite_scan(text)
            lolbas = scan_lolbas(text)
            risk = risk_score(mitre, yara, iocs)
            partial = {"iocs": iocs, "mitre": mitre, "yara": yara, "lolbas": lolbas, "risk": risk}
            yield _sse("partial", partial)
        except Exception as e:
            yield _sse("error", {"phase": "extract", "error": str(e)})
            return

        # 2) TI hits (fast — local DB)
        try:
            yield _sse("status", {"phase": "ti_hits", "message": "Cross-referencing local Threat-Intel DB…"})
            ti_hits = await _lookup_ti_hits(iocs)
            yield _sse("ti_hits", ti_hits)
        except Exception as e:
            ti_hits = []
            yield _sse("error", {"phase": "ti_hits", "error": str(e)})

        # 3) Parallelize OSINT + AI describe/verdict, streaming a heartbeat while we wait
        yield _sse("status", {"phase": "enrich_and_ai", "message": "Running OSINT enrichment + AI analysis in parallel…"})

        async def _run_osint():
            if not body.enrich_osint:
                return None
            keys = await load_osint_keys()
            return await enrich_iocs(iocs, keys)

        async def _run_ai():
            if not (body.use_ai_verdict or body.describe):
                return None
            return await _ai_describe_and_verdict(
                body.input, body.output or "", iocs, mitre, yara, {},
                lolbas=lolbas,
                want_verdict=body.use_ai_verdict, want_describe=body.describe,
            )

        osint_task = asyncio.create_task(_run_osint())
        ai_task = asyncio.create_task(_run_ai())
        pending = {osint_task, ai_task}

        # Heartbeat every 10s to keep Cloudflare / reverse-proxy connection alive
        elapsed = 0
        osint_data: Optional[Dict[str, Any]] = None
        ai_bundle: Optional[Dict[str, Any]] = None
        while pending:
            done, pending = await asyncio.wait(pending, timeout=10.0, return_when=asyncio.FIRST_COMPLETED)
            elapsed += 10
            if not done:
                yield _sse("heartbeat", {"elapsed_s": elapsed, "phase": "waiting"})
                # SSE comment as extra keep-alive nudge for proxies
                yield b": keep-alive\n\n"
                continue
            for t in done:
                try:
                    r = t.result()
                except Exception as e:
                    if t is osint_task:
                        osint_data = {"error": str(e)}
                        yield _sse("error", {"phase": "osint", "error": str(e)})
                    else:
                        ai_bundle = None
                        yield _sse("error", {"phase": "ai", "error": str(e)})
                    continue
                if t is osint_task:
                    osint_data = r
                    yield _sse("osint", osint_data or {})
                else:
                    ai_bundle = r
                    if ai_bundle:
                        if body.use_ai_verdict and ai_bundle.get("verdict") is not None:
                            yield _sse("ai_verdict", ai_bundle.get("verdict"))
                        if body.describe and ai_bundle.get("description") is not None:
                            yield _sse("description", ai_bundle.get("description"))

        ai_verdict = ai_bundle.get("verdict") if (ai_bundle and body.use_ai_verdict) else None
        description = ai_bundle.get("description") if (ai_bundle and body.describe) else None

        # Merge AI-derived MITRE with heuristics
        merged_mitre = list(mitre)
        if description and isinstance(description, dict):
            ai_mitre = description.get("mitre_techniques") or []
            seen_ids = {m["id"] for m in merged_mitre}
            for m in ai_mitre:
                if isinstance(m, dict) and m.get("id") and m["id"] not in seen_ids:
                    merged_mitre.append({
                        "id": m["id"], "technique": m.get("technique", ""),
                        "tactic": m.get("tactic", ""), "evidence": m.get("evidence", ""),
                        "source": "ai",
                    })
                    seen_ids.add(m["id"])
            for m in merged_mitre:
                m.setdefault("source", "heuristic")

        final = {
            "iocs": iocs, "mitre": merged_mitre, "yara": yara, "lolbas": lolbas, "risk": risk,
            "osint": osint_data, "ti_hits": ti_hits,
            "ai_verdict": ai_verdict, "description": description,
        }
        yield _sse("result", final)
        yield _sse("done", {"ok": True})

    return StreamingResponse(
        gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",  # disable nginx buffering
            "Connection": "keep-alive",
        },
    )


# =============================================================================
# Async job pipeline — /api/analyze/async + /api/analyze/status/{job_id}
# Bypasses reverse-proxy hard timeouts for long AI runs by decoupling the
# background computation from client polling. Jobs live in MongoDB
# (`analyze_jobs` collection) with a TTL index so they survive backend restarts
# and work correctly across multiple replicas.
# =============================================================================
import uuid

_JOB_TTL_SEC = 60 * 15  # 15 minutes
_ANALYZE_JOBS_INDEX_READY = False


async def _ensure_jobs_indexes():
    """Idempotent — creates a TTL index on `created_at` for auto-cleanup."""
    global _ANALYZE_JOBS_INDEX_READY
    if _ANALYZE_JOBS_INDEX_READY:
        return
    try:
        await db.analyze_jobs.create_index(
            "created_at", expireAfterSeconds=_JOB_TTL_SEC, name="ttl_created_at"
        )
        _ANALYZE_JOBS_INDEX_READY = True
    except Exception as e:
        log.warning("analyze_jobs TTL index create failed: %s", e)


async def _job_get(job_id: str) -> Optional[Dict[str, Any]]:
    doc = await db.analyze_jobs.find_one({"_id": job_id}, {"created_at": 0})
    return doc


async def _job_set(job_id: str, updates: Dict[str, Any]) -> None:
    """Merge `updates` into the job doc. `updated_at` is refreshed every write."""
    updates = {k: v for k, v in updates.items() if not k.startswith("_")}
    updates["updated_at"] = datetime.now(timezone.utc)
    await db.analyze_jobs.update_one({"_id": job_id}, {"$set": updates}, upsert=False)


async def _job_push_error(job_id: str, phase: str, err: str) -> None:
    await db.analyze_jobs.update_one(
        {"_id": job_id},
        {"$push": {"errors": {"phase": phase, "error": err}},
         "$set": {"updated_at": datetime.now(timezone.utc)}},
    )


async def _run_analysis_job(job_id: str, body: AnalyzeIn):
    """Background worker — persists phase-by-phase progress to MongoDB."""
    try:
        # Phase A — fast extraction (+ custom detection rules from Model Studio)
        text = (body.output or "") + "\n" + body.input
        iocs = extract_iocs(text)
        mitre = mitre_map(text)
        yara = yara_lite_scan(text)
        lolbas = scan_lolbas(text)
        try:
            custom_rules = await ms._load_active_rules(db)
            custom_hits = ms.scan_custom_rules(text, custom_rules)
            for h in custom_hits:
                if h.get("model_id"):
                    await ms.increment_usage(db, h["model_id"])
            lolbas = lolbas + custom_hits
        except Exception as e:
            log.warning("custom detection rules failed: %s", e)
        risk = risk_score(mitre, yara, iocs)
        await _job_set(job_id, {
            "iocs": iocs, "mitre": mitre, "yara": yara, "lolbas": lolbas, "risk": risk,
            "phase": "ti_hits", "progress": 15,
        })

        # Phase B — TI hits
        ti_hits = await _lookup_ti_hits(iocs)
        await _job_set(job_id, {
            "ti_hits": ti_hits, "phase": "enrich_and_ai", "progress": 25,
        })

        # Phase C — OSINT + AI in parallel (with optional persona/provider + all enabled playbooks)
        persona = await ms.get_persona(db, body.persona_id) if body.persona_id else None
        provider = await ms.get_provider(db, body.provider_id)
        playbook_block, playbooks_used = await ms.compose_playbook_prompt_with_meta(db, target="ai")
        # Persist the snapshot so /analyze/{job_id}/feedback can attribute votes.
        await _job_set(job_id, {"playbooks_used": playbooks_used})
        if persona and persona.get("id"):
            await ms.increment_usage(db, persona["id"])
        if provider and provider.get("id") and body.provider_id:
            await ms.increment_usage(db, provider["id"])

        async def _run_osint():
            if not body.enrich_osint:
                return None
            keys = await load_osint_keys()
            return await enrich_iocs(iocs, keys)

        async def _run_ai():
            if not (body.use_ai_verdict or body.describe):
                return None
            return await _ai_describe_and_verdict(
                body.input, body.output or "", iocs, mitre, yara, {},
                lolbas=lolbas,
                want_verdict=body.use_ai_verdict, want_describe=body.describe,
                persona=persona, provider=provider,
                playbook=playbook_block,
            )

        osint_task = asyncio.create_task(_run_osint())
        ai_task = asyncio.create_task(_run_ai())
        pending = {osint_task, ai_task}
        osint_data: Optional[Dict[str, Any]] = None
        ai_bundle: Optional[Dict[str, Any]] = None
        start = datetime.now(timezone.utc).timestamp()

        while pending:
            done, pending = await asyncio.wait(pending, timeout=2.0, return_when=asyncio.FIRST_COMPLETED)
            elapsed_s = int(datetime.now(timezone.utc).timestamp() - start)
            heartbeat: Dict[str, Any] = {"elapsed_s": elapsed_s}
            for t in done:
                try:
                    r = t.result()
                except Exception as e:
                    if t is osint_task:
                        osint_data = {"error": str(e)}
                        heartbeat["osint"] = osint_data
                        await _job_push_error(job_id, "osint", str(e))
                    else:
                        ai_bundle = None
                        await _job_push_error(job_id, "ai", str(e))
                    continue
                if t is osint_task:
                    osint_data = r
                    heartbeat["osint"] = osint_data
                    heartbeat["progress"] = 45
                else:
                    ai_bundle = r
                    if ai_bundle:
                        if body.use_ai_verdict and ai_bundle.get("verdict") is not None:
                            heartbeat["ai_verdict"] = ai_bundle.get("verdict")
                        if body.describe and ai_bundle.get("description") is not None:
                            heartbeat["description"] = ai_bundle.get("description")
                    heartbeat["progress"] = 90
            await _job_set(job_id, heartbeat)

        # Phase D — merge & finalize
        merged_mitre = list(mitre)
        description = ai_bundle.get("description") if ai_bundle else None
        if description and isinstance(description, dict):
            ai_mitre = description.get("mitre_techniques") or []
            seen_ids = {m["id"] for m in merged_mitre}
            for m in ai_mitre:
                if isinstance(m, dict) and m.get("id") and m["id"] not in seen_ids:
                    merged_mitre.append({
                        "id": m["id"], "technique": m.get("technique", ""),
                        "tactic": m.get("tactic", ""), "evidence": m.get("evidence", ""),
                        "source": "ai",
                    })
                    seen_ids.add(m["id"])
            for m in merged_mitre:
                m.setdefault("source", "heuristic")
        await _job_set(job_id, {
            "mitre": merged_mitre,
            "status": "done",
            "progress": 100,
            "phase": "complete",
        })
    except Exception as e:
        log.exception("analysis job %s failed", job_id)
        await _job_set(job_id, {"status": "error", "error": str(e), "phase": "error"})


@api.post("/analyze/async")
async def analyze_async(body: AnalyzeIn, user=Depends(get_current_user)):
    """Kick off an analysis job. Returns immediately with a job_id."""
    await _ensure_jobs_indexes()
    job_id = uuid.uuid4().hex
    now = datetime.now(timezone.utc)
    await db.analyze_jobs.insert_one({
        "_id": job_id,
        "status": "running",
        "phase": "extract",
        "progress": 5,
        "created_at": now,      # TTL-indexed
        "updated_at": now,
        "requested_by": user.get("email"),
    })
    asyncio.create_task(_run_analysis_job(job_id, body))
    return {"job_id": job_id, "status": "running"}


@api.get("/analyze/status/{job_id}")
async def analyze_status(job_id: str, user=Depends(get_current_user)):
    doc = await _job_get(job_id)
    if not doc:
        raise HTTPException(status_code=404, detail="job not found or expired")
    # Rename Mongo `_id` → `job_id` for the API client
    doc["job_id"] = doc.pop("_id")
    # Serialize datetime → ISO
    for k in ("updated_at",):
        if isinstance(doc.get(k), datetime):
            doc[k] = doc[k].isoformat()
    return doc


class PlaybookFeedbackIn(BaseModel):
    vote: str = Field(..., description="'up', 'down' or 'none' (to retract)")
    reason: Optional[str] = None


@api.post("/analyze/{job_id}/feedback")
async def analyze_feedback(job_id: str, body: PlaybookFeedbackIn,
                            user=Depends(get_current_user)):
    """Record a 👍/👎 vote on the AI investigation attached to `job_id`.

    Toggling is allowed — the previous vote (if any) is reversed on all playbooks
    that were used before the new one is applied. A full audit trail of every
    vote change is appended to `playbook_votes.history`.
    """
    job = await _job_get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="job not found or expired")
    playbooks_used = job.get("playbooks_used") or []
    if not playbooks_used:
        raise HTTPException(status_code=400,
                            detail="this investigation did not apply any playbook — nothing to feedback on")
    try:
        result = await ms.record_playbook_vote(
            db, job_id, user.get("email") or "anonymous",
            playbooks_used, body.vote, body.reason,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    result["playbooks_used"] = playbooks_used
    return result


@api.get("/analyze/{job_id}/feedback")
async def analyze_feedback_get(job_id: str, user=Depends(get_current_user)):
    """Return the current user's vote (if any) + the playbooks attributed to this job."""
    job = await _job_get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="job not found or expired")
    playbooks_used = job.get("playbooks_used") or []
    vote = await ms.get_vote_for_job(db, job_id, user.get("email") or "anonymous")
    return {"job_id": job_id, "playbooks_used": playbooks_used, **vote}


@api.get("/admin/playbooks/{playbook_id}/votes")
async def playbook_votes(playbook_id: str, limit: int = 50, user=Depends(require_admin)):
    """Full audit trail of the last `limit` votes attributed to this playbook."""
    return {
        "playbook_id": playbook_id,
        "votes": await ms.list_playbook_votes(db, playbook_id, limit=limit),
    }


async def _ai_describe_and_verdict(inp, out, iocs, mitre, yara, osint, want_verdict, want_describe, lolbas=None,
                                    persona: Optional[Dict[str, Any]] = None,
                                    provider: Optional[Dict[str, Any]] = None,
                                    playbook: str = ""):
    """Single LLM call producing rich narrative description + verdict JSON.

    Optional `persona`   overrides the default system prompt.
    Optional `provider` swaps the LLM provider/model (Claude, GPT, Gemini).
    Optional `playbook` — free-form analyst guidance appended to the system prompt.
    """
    parts = []
    if want_describe:
        parts.append(
            '"description": {\n'
            '  "summary": "2-3 sentence executive summary of what the decoded script/command does",\n'
            '  "malware_family": {\n'
            '     "name": "concrete family/tooling name if identifiable (e.g. Cobalt Strike, AsyncRAT, Emotet, Nanocore, XORDDoS, Empire, custom Python loader) or null if unknown",\n'
            '     "confidence": "low|medium|high",\n'
            '     "rationale": "why this family — cite specific TTPs, string patterns, key material, structure, or matches to public reports"\n'
            '  },\n'
            '  "mitre_techniques": [\n'
            '     {"id": "Txxxx or Txxxx.xxx", "technique": "name", "tactic": "MITRE tactic (Execution|Defense Evasion|...)", "evidence": "specific line/token in the decoded output that supports this mapping"}\n'
            '  ],\n'
            '  "attack_chain": [\n'
            '     {\n'
            '        "step": 1,\n'
            '        "title": "SHORT TECHNICAL LABEL (e.g. WRAPPER DEOBFUSCATION, CONTEXT ADJUSTMENT, STEALTH STAGER LOAD, ROLLING XOR DECRYPTION, FILELESS MEMORY INJECTION, C2 BEACONING, PERSISTENCE INSTALL, SHADOW COPY DESTROY)",\n'
            '        "summary": "2-3 sentence plain-English narrative of what THIS step does. Be specific: name the API/function invoked, files touched, algorithms used.",\n'
            '        "technical_detail": "optional: paste concrete artifact from the payload (e.g. hex key, filename, URL, IP, argv) that supports this step. null if none.",\n'
            '        "kind": "ingestion|deobfuscation|context|filesystem|network|crypto|execution|persistence|discovery|c2|impact"\n'
            '     }\n'
            '  ],\n'
            '  "entity_graph": {\n'
            '     "nodes": [\n'
            '        {"id": "e1", "label": "ADVERSARY OBJECTIVE ONLY — describe WHAT the attacker is achieving on the victim, not the script mechanic. Good examples: \'Initial Foothold via Malicious Document\', \'Encrypted Stager Delivery\', \'Fileless In-Memory Execution\', \'Defense Evasion via Extension Masquerade\', \'C2 Beacon Establishment\', \'Credential Harvesting\', \'Data Exfiltration\', \'Shadow Copy Destruction (Anti-Recovery)\'. BAD examples that must NEVER appear: \'python.exe\', \'instructions.docx\', \'XOR Key 4fab...\', \'base64.b64decode()\', \'os.chdir\', \'exec()\', filenames, function calls, hex strings, variable names.", "type": "action|ip|url|user|device", "tactic": "MITRE ATT&CK tactic — REQUIRED — Reconnaissance|Resource Development|Initial Access|Execution|Persistence|Privilege Escalation|Defense Evasion|Credential Access|Discovery|Lateral Movement|Collection|Command and Control|Exfiltration|Impact", "malicious": true, "note": "optional 1-line summary of the attacker\'s intent at this step"}\n'
            '     ],\n'
            '     "edges": [\n'
            '        {"from": "e1", "to": "e2", "label": "attacker progression verb (Enables, Escalates To, Leverages For, Feeds Into, Culminates In, Establishes, Exfiltrates To). NEVER script mechanic verbs like Reads, Loads, Decrypts, Parent Of."}\n'
            '     ]\n'
            '  },\n'
            '  "flow_graph": {\n'
            '     "nodes": [\n'
            '        {"id": "n1", "label": "short verb-phrase e.g. \'chdir to python.exe folder\'", "kind": "start|filesystem|network|crypto|execution|persistence|discovery|c2|impact|end"}\n'
            '     ],\n'
            '     "edges": [\n'
            '        {"from": "n1", "to": "n2", "label": "optional: describes transition / data flow"}\n'
            '     ]\n'
            '  },\n'
            '  "behavior": ["bullet points describing each behavior observed"],\n'
            '  "ioc_narrative": "1-2 paragraph narrative discussing extracted IOCs, referencing OSINT enrichment where present (VT verdict, AbuseIPDB score, Shodan ports, TI-hits, geolocation). Be specific with values.",\n'
            '  "attribution_hints": "any hints about actor / campaign / open-source tooling / commodity vs targeted",\n'
            '  "recommended_actions": ["array of concrete containment / IR actions"]\n'
            '}'
        )
    if want_verdict:
        parts.append(
            '"verdict": {\n'
            '  "verdict": "Malicious|Suspicious|Benign",\n'
            '  "confidence": 0-100,\n'
            '  "summary": "1-2 sentence rationale — always mention malware family if identified",\n'
            '  "key_findings": ["short strings"],\n'
            '  "recommended_actions": ["short strings"]\n'
            '}'
        )
    schema = "{\n" + ",\n".join(parts) + "\n}"

    default_system = (
        "You are a senior DFIR analyst reviewing a decoded payload. "
        "Write like an incident-report analyst: precise, factual, technical, cite specific IOC values / OSINT results / TI hits.\n"
        "For malware_family: only claim a family if there is strong evidence (unique strings, C2 patterns, packer, algorithm signatures, or matches to VT/OTX threat labels).\n"
        "For mitre_techniques: derive from the DECODED BEHAVIOR, not the outer wrapper. For each technique cite the specific evidence in the decoded output.\n"
        "For attack_chain: model 3-8 sequential steps that describe the ATTACK CHAIN (what the malware does step by step). Each step should have a strong technical title (short caps phrase) + 2-3 sentence plain-English summary + optional technical_detail (hex keys, filenames, URLs cited verbatim from the payload). Order MUST be causal. Use kinds: ingestion|deobfuscation|context|filesystem|network|crypto|execution|persistence|discovery|c2|impact.\n"
        "For entity_graph: this is a TACTICAL ATTACK CHAIN describing what THE ATTACKER is achieving on the victim system — NOT a technical decoder trace of the script's operations. Extract 5-10 ATTACKER GOALS/OUTCOMES, each mapped to a MITRE ATT&CK tactic. Frame every node from the adversary's perspective (their objectives on the target).\n"
        "  STRICT RULES:\n"
        "  - Do NOT create nodes for filenames, functions, APIs, variables, hex keys, or script internals (e.g. 'python.exe', 'instructions.docx', 'XOR Key 4fab…', 'base64.b64decode', 'os.chdir'). These are HOW the attacker operates, not WHAT they achieve.\n"
        "  - DO create nodes for adversary objectives (e.g. 'Initial Foothold via Malicious Document', 'Encrypted Stager Delivery', 'In-Memory Fileless Execution', 'Defense Evasion via Extension Masquerade', 'C2 Establishment', 'Credential Harvesting Attempt', 'Data Exfiltration', 'Shadow Copy Destruction (Anti-Recovery)').\n"
        "  - Node type should usually be 'action' (attacker action/objective) except for concrete external entities such as C2 IPs/URLs (type ip/url) or targeted victim entities (type user/device).\n"
        "  - Every node MUST carry a MITRE ATT&CK tactic. Order the graph as a kill-chain timeline (Initial Access → Execution → Defense Evasion → Persistence → Credential Access → Discovery → Command and Control → Collection → Exfiltration → Impact).\n"
        "  - Edges must express attacker progression (e.g. 'Enables', 'Escalates To', 'Leverages For', 'Feeds Into', 'Culminates In') — NOT script data-flow like 'Reads', 'Loads', 'Decrypts'.\n"
        "  - Include the malicious flag on adversary-controlled nodes.\n"
        "  Example: for a Python XOR loader dropping a stager, DO NOT list 'python.exe','base64','XOR key','instructions.docx'. INSTEAD produce nodes like: {label:'Initial Foothold (Signed LOLBin Abuse)', tactic:'Initial Access'} → {label:'Encrypted Stager Retrieval', tactic:'Defense Evasion'} → {label:'Rolling-Key Deobfuscation of Second Stage', tactic:'Defense Evasion'} → {label:'Fileless In-Memory Payload Execution', tactic:'Execution'} → {label:'C2 Beacon Establishment', tactic:'Command and Control'}.\n"
        "For flow_graph: additionally produce a compact node/edge structure for visualization — 4-10 nodes.\n"
        "Return STRICT JSON only with the keys shown in the schema. No markdown, no prose outside JSON."
    )
    # Custom persona (if selected) overrides the default system prompt. When the
    # persona is a free-form narrative persona (like the "Static Malware Deobfuscation
    # Engine"), we still append a "return JSON with this schema" directive so the
    # /analyze pipeline keeps working with structured output. Otherwise the AI tab
    # would receive unparseable prose and every downstream card would fall back.
    if persona and (persona.get("config") or {}).get("system_prompt"):
        persona_prompt = persona["config"]["system_prompt"].strip()
        system = (
            persona_prompt
            + "\n\nIMPORTANT — OUTPUT CONTRACT (in addition to your normal analysis):\n"
            + "Return your final analysis as STRICT JSON only, matching the schema below. "
            + "Fold your persona-specific findings into the `summary`, `attack_chain`, `behavior`, and `ioc_narrative` fields. "
            + "No markdown, no prose outside JSON."
        )
    else:
        system = default_system

    # Append the org-wide analyst playbook (playbook kind from Model Studio)
    if playbook:
        system = system + "\n" + playbook

    # LLM provider / model — default to Claude Sonnet 4.5
    llm_provider = ((provider or {}).get("config") or {}).get("provider") or "anthropic"
    llm_model = ((provider or {}).get("config") or {}).get("model") or "claude-sonnet-4-5-20250929"

    prompt = (
        f"SCHEMA:\n{schema}\n\n"
        f"RAW INPUT:\n{inp[:3500]}\n\n"
        f"DECODED OUTPUT:\n{out[:3500]}\n\n"
        f"EXTRACTED IOCs:\n{json.dumps(iocs)[:2000]}\n\n"
        f"HEURISTIC MITRE (from wrapper text):\n{json.dumps(mitre)[:1200]}\n\n"
        f"HEURISTIC YARA:\n{json.dumps(yara)[:1200]}\n\n"
        f"LOLBAS MATCHES:\n{json.dumps(lolbas or [])[:1500]}\n\n"
        f"OSINT ENRICHMENT:\n{json.dumps(osint)[:5000]}\n\n"
        "Return only JSON."
    )
    return await _llm_json(
        "describe-" + str(datetime.now(timezone.utc).timestamp()),
        system, prompt,
        provider=llm_provider, model=llm_model,
    )


# =============================================================================
# Endpoints — AI (Auto Decode / Auto Investigate / Troubleshoot)
# =============================================================================
_OP_IDS = sorted(OPERATIONS.keys())


@api.post("/ai/auto-decode")
async def ai_auto_decode(body: AutoIn, user=Depends(get_current_user)):
    """AI Decode with strict anti-hallucination guardrails for SOC use.

    Pipeline:
      1. Ask the LLM for a decoding recipe (JSON) — its ARGS included.
      2. Run BOTH the AI plan AND the deterministic magic decoder in parallel.
      3. Score each candidate (readability + shellcode prologue + IOC density).
      4. Pick the higher-scoring winner. If BOTH are below the quality floor,
         STOP GRACEFULLY and return `"No further deterministic decoding
         possible"` with a confidence score + the exact ops attempted.

    The output is NEVER silently corrupted — we return the winning engine, a
    confidence 0–100, a `stopped_gracefully` flag, and per-step previews.
    """
    system = (
        "You are an expert malware analyst using a CyberChef-like tool. "
        "Given an obfuscated / encoded payload, produce a JSON recipe of operations that will fully decode it.\n"
        f"AVAILABLE OPERATION IDS: {_OP_IDS}\n"
        "Return STRICT JSON only with keys: reasoning (short string), steps (array of {op, args}).\n"
        "For XOR ops YOU MUST include the key in args, e.g. {\"op\":\"xor\",\"args\":{\"key\":\"0x23\"}}. "
        "For gzip/zlib/base64 chained forms, list each step separately. "
        "Only use ids from AVAILABLE OPERATION IDS. Max 8 steps."
    )
    prompt = f"PAYLOAD:\n{body.input[:4000]}\n\nReturn only JSON."
    try:
        plan = await _llm_json("autodecode-" + str(datetime.now(timezone.utc).timestamp()), system, prompt)
    except HTTPException:
        plan = {"reasoning": "AI unavailable — falling back to deterministic decoder.", "steps": []}

    # Build AI recipe
    ai_steps: List[RecipeStep] = []
    for s in (plan.get("steps") or [])[:8]:
        if s.get("op") in OPERATIONS:
            ai_steps.append(RecipeStep(op=s["op"], args=s.get("args") or {}))

    # Execute AI plan (if any) + magic_decode in parallel
    async def _run_ai_plan() -> Dict[str, Any]:
        if not ai_steps:
            return {"engine": "ai", "output": "", "recipe": [], "errors": ["no valid steps proposed"]}
        r = await run_recipe(RunRecipeIn(input=body.input, steps=ai_steps), user=user)
        return {
            "engine": "ai",
            "output": r.output or "",
            "recipe": [s.model_dump() for s in ai_steps],
            "steps_output": r.steps_output,
            "detected_type": r.detected_type,
            "errors": r.errors,
        }

    def _run_magic() -> Dict[str, Any]:
        m = magic_decode(body.input, max_depth=6, max_branches=5, top_n=3)
        top = (m.get("top_results") or [{}])[0]
        return {
            "engine": "magic",
            "output": top.get("output", ""),
            "recipe": [{"op": c["op"], "args": c.get("args") or {}}
                       for c in (top.get("chain") or [])],
            "steps_output": top.get("chain") or [],
            "detected_type": m.get("detected_type"),
            "errors": [],
            "is_shellcode": top.get("is_shellcode", False),
            "score": (top.get("score_breakdown") or {}).get("score", 0.0),
        }

    ai_task = asyncio.create_task(_run_ai_plan())
    magic_result = _run_magic()
    ai_result = await ai_task

    # Quality scoring — the anti-hallucination guard
    def _quality_score(res: Dict[str, Any]) -> Dict[str, Any]:
        out = res.get("output") or ""
        if not out:
            return {"score": 0.0, "printable_ratio": 0.0, "shellcode": False,
                    "iocs_found": 0, "reasons": ["empty output"]}
        printable = sum(1 for c in out if 32 <= ord(c) < 127 or ord(c) in (9, 10, 13))
        pr = printable / max(1, len(out))
        reasons: List[str] = []
        score = pr * 0.4  # 40% weight on readability

        # Shellcode prologue detection — strong positive signal
        try:
            from shellcode_analyzer import starts_with_known_prologue
            raw = out.encode("latin-1") if all(ord(c) < 256 for c in out) \
                                        else out.encode("utf-8", errors="replace")
            is_sc = starts_with_known_prologue(raw)
            if is_sc:
                score += 0.35
                reasons.append("known shellcode prologue detected (+0.35)")
        except Exception:
            is_sc = False

        # IOC extraction — any URLs/IPs/hashes recovered is a strong signal
        try:
            iocs = extract_iocs_from_text(out)
            n_iocs = sum(len(v) if isinstance(v, list) else 0 for v in iocs.values())
            if n_iocs:
                score += min(0.20, 0.05 * n_iocs)
                reasons.append(f"{n_iocs} IOC(s) recovered")
        except Exception:
            n_iocs = 0

        # Known script markers
        markers = ("IEX", "Invoke-Expression", "DownloadString", "New-Object",
                   "FromBase64String", "http://", "https://", "MZ", "PE", "ELF",
                   "cmd.exe", "powershell", "$env:")
        matched_markers = [m for m in markers if m in out]
        if matched_markers:
            score += min(0.15, 0.05 * len(matched_markers))
            reasons.append(f"script markers: {matched_markers[:3]}")

        # Bonus: readability high enough that this looks like real text
        if pr >= 0.90 and not is_sc:
            score += 0.10
            reasons.append("clean printable (+0.10)")

        return {
            "score": round(min(1.0, score), 3),
            "printable_ratio": round(pr, 3),
            "shellcode": is_sc,
            "iocs_found": n_iocs,
            "reasons": reasons,
        }

    ai_q = _quality_score(ai_result)
    mg_q = _quality_score(magic_result)

    # Quality floor for SOC use: we WILL NOT return an output below this.
    QUALITY_FLOOR = 0.35

    # HARD RULE: if one engine reached shellcode terminal state and the other
    # did NOT, the shellcode-reaching engine WINS unconditionally. This
    # prevents the AI from being rewarded for adding gratuitous "Extract
    # Strings" / "clean up" steps that mangle raw shellcode into fragmented
    # ASCII (which would otherwise score higher on printable_ratio).
    if ai_q["shellcode"] and not mg_q["shellcode"]:
        winner, wq, loser, lq = ai_result, ai_q, magic_result, mg_q
    elif mg_q["shellcode"] and not ai_q["shellcode"]:
        winner, wq, loser, lq = magic_result, mg_q, ai_result, ai_q
    elif ai_q["score"] >= mg_q["score"]:
        winner, wq, loser, lq = ai_result, ai_q, magic_result, mg_q
    else:
        winner, wq, loser, lq = magic_result, mg_q, ai_result, ai_q

    stopped_gracefully = wq["score"] < QUALITY_FLOOR

    response = {
        "reasoning": plan.get("reasoning", ""),
        "recipe": winner["recipe"],
        "output": winner["output"] if not stopped_gracefully else "",
        "steps_output": winner.get("steps_output") or [],
        "detected_type": winner.get("detected_type"),
        "errors": winner.get("errors") or [],
        # Anti-hallucination guardrail metadata
        "winner_engine": winner["engine"],
        "confidence": int(round(wq["score"] * 100)),
        "quality_reasons": wq["reasons"],
        "stopped_gracefully": stopped_gracefully,
        "graceful_message": (
            "No further deterministic decoding possible. "
            f"Best attempt via '{winner['engine']}' engine scored {int(round(wq['score']*100))}/100 "
            f"(readability {int(wq['printable_ratio']*100)}%, "
            f"shellcode={wq['shellcode']}, IOCs={wq['iocs_found']}). "
            "The payload may already be plaintext, use a key/format not yet supported, "
            "or be intentionally corrupted."
        ) if stopped_gracefully else "",
        "alternate": {
            "engine": loser["engine"],
            "confidence": int(round(lq["score"] * 100)),
            "recipe": loser.get("recipe") or [],
        },
    }
    return response


def extract_iocs_from_text(text: str) -> Dict[str, List[str]]:
    """Lightweight IOC extraction reused by the quality gate."""
    from command_analyzer import extract_iocs as _ex
    r = _ex(text or "")
    # Flatten nested hashes dict for counting convenience
    flat = {"urls": r.get("urls") or [], "ips": r.get("ips") or [],
            "regkeys": r.get("regkeys") or [], "file_paths": r.get("file_paths") or []}
    h = r.get("hashes") or {}
    flat["hashes"] = (h.get("md5") or []) + (h.get("sha1") or []) + (h.get("sha256") or [])
    return flat


@api.post("/ai/auto-investigate")
async def ai_auto_investigate(body: AutoIn, user=Depends(get_current_user)):
    """Auto Decode + full Analyze (OSINT + AI describe + AI verdict) — optimized.

    Strategy for speed:
      1. Run deterministic smart-decode FIRST (instant, no AI wait).
      2. If smart decoder found nothing, fall back to AI decode.
      3. Run OSINT enrichment and AI describe/verdict IN PARALLEL against the decoded output.
    """
    # 1) fast deterministic decode
    det = smart_decode(body.input)
    if det["steps"]:
        steps = [RecipeStep(op=x["op"], args=x.get("args", {})) for x in det["steps"]]
        reasoning = "Deterministic smart decoder chained: " + " → ".join(s.op for s in steps)
    else:
        # 2) fall back to AI-planned decode
        dec = await ai_auto_decode(body, user=user)
        steps = [RecipeStep(op=s["op"], args=s.get("args", {})) for s in dec["recipe"]]
        reasoning = dec.get("reasoning", "")

    exec_result = await run_recipe(RunRecipeIn(input=body.input, steps=steps), user=user)
    decoded_output = exec_result.output

    # 3) analyze (with OSINT + AI describe + AI verdict + LOLBAS + TI-hits)
    analysis = await analyze(AnalyzeIn(
        input=body.input, output=decoded_output,
        use_ai_verdict=True, describe=True, enrich_osint=True,
    ), user=user)

    return {
        "reasoning": reasoning,
        "recipe": [s.model_dump() for s in steps],
        "output": decoded_output,
        "steps_output": exec_result.steps_output,
        "detected_type": exec_result.detected_type,
        "errors": exec_result.errors,
        "analysis": analysis,
    }


@api.post("/ai/troubleshoot")
async def ai_troubleshoot(body: TroubleshootIn, user=Depends(get_current_user)):
    system = (
        "You are a DFIR analyst helping troubleshoot a stuck decoding recipe. "
        "Given the input, the recipe applied, and any error, explain what went wrong (1-3 sentences) "
        "and propose a fixed recipe.\n"
        f"AVAILABLE OPERATION IDS: {_OP_IDS}\n"
        "Return STRICT JSON: {diagnosis: string, suggested_steps: [{op, args}]}. Max 8 steps."
    )
    prompt = (
        f"INPUT:\n{body.input[:3000]}\n\n"
        f"CURRENT RECIPE: {json.dumps([s.model_dump() for s in body.steps])}\n\n"
        f"ERROR: {body.error or 'no error - output looks wrong'}\n\nReturn only JSON."
    )
    result = await _llm_json("troubleshoot-" + str(datetime.now(timezone.utc).timestamp()), system, prompt)
    fixed = [{"op": s["op"], "args": s.get("args") or {}} for s in (result.get("suggested_steps") or [])[:8] if s.get("op") in OPERATIONS]
    return {"diagnosis": result.get("diagnosis", ""), "suggested_steps": fixed}


# =============================================================================
# Share / Report
# =============================================================================
@api.post("/share")
async def create_share(body: ShareIn, user=Depends(get_current_user)):
    payload = json.dumps({"input": body.input, "steps": [s.model_dump() for s in body.steps]}).encode("utf-8")
    token = base64.urlsafe_b64encode(payload).decode("utf-8").rstrip("=")
    await db.shares.insert_one({
        "token": token, "input_len": len(body.input),
        "steps": [s.model_dump() for s in body.steps],
        "created_by": user["email"],
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"token": token}


@api.get("/share/{token}")
async def get_share(token: str):
    padded = token + "=" * (-len(token) % 4)
    try:
        return json.loads(base64.urlsafe_b64decode(padded).decode("utf-8"))
    except Exception:
        raise HTTPException(status_code=404, detail="Invalid share token")


@api.post("/report")
async def build_report(body: AnalyzeIn, user=Depends(get_current_user)):
    """JSON with both text + html body (for backward compat / preview)."""
    ctx = await _analysis_context(body, user)
    ts = ctx["ts"]
    txt = _render_text_report(user, ts, body, ctx["risk"], ctx["mitre"], ctx["yara"], ctx["lolbas"], ctx["iocs"], ctx["ti_hits"], ctx["osint"], ctx["description"], ctx["verdict"])
    html = _render_html_report(user, ts, body, ctx["risk"], ctx["mitre"], ctx["yara"], ctx["lolbas"], ctx["iocs"], ctx["ti_hits"], ctx["osint"], ctx["description"], ctx["verdict"])
    return {
        "report": txt,
        "html": html,
        "filename": f"nivxray_report_{int(datetime.now().timestamp())}.txt",
        "filename_html": f"nivxray_report_{int(datetime.now().timestamp())}.html",
    }


@api.post("/report/{fmt}")
async def build_report_fmt(fmt: str, body: AnalyzeIn, user=Depends(get_current_user)):
    """Download report as txt / html / csv / docx / pdf."""
    fmt = fmt.lower()
    if fmt not in ("txt", "html", "csv", "docx", "pdf"):
        raise HTTPException(status_code=400, detail="format must be one of txt|html|csv|docx|pdf")
    ctx = await _analysis_context(body, user)
    stem = f"nivxray_report_{int(datetime.now().timestamp())}"
    ts = ctx["ts"]

    if fmt == "txt":
        payload = _render_text_report(user, ts, body, ctx["risk"], ctx["mitre"], ctx["yara"], ctx["lolbas"], ctx["iocs"], ctx["ti_hits"], ctx["osint"], ctx["description"], ctx["verdict"]).encode("utf-8")
        return _download(payload, f"{stem}.txt", "text/plain; charset=utf-8")

    if fmt == "html":
        payload = _render_html_report(user, ts, body, ctx["risk"], ctx["mitre"], ctx["yara"], ctx["lolbas"], ctx["iocs"], ctx["ti_hits"], ctx["osint"], ctx["description"], ctx["verdict"]).encode("utf-8")
        return _download(payload, f"{stem}.html", "text/html; charset=utf-8")

    if fmt == "csv":
        payload = _render_csv_report(user, ts, body, ctx).encode("utf-8")
        return _download(payload, f"{stem}.csv", "text/csv; charset=utf-8")

    if fmt == "docx":
        payload = _render_docx_report(user, ts, body, ctx)
        return _download(payload, f"{stem}.docx",
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if fmt == "pdf":
        html_body = _render_html_report(user, ts, body, ctx["risk"], ctx["mitre"], ctx["yara"], ctx["lolbas"], ctx["iocs"], ctx["ti_hits"], ctx["osint"], ctx["description"], ctx["verdict"])
        payload = _render_pdf_from_html(html_body)
        return _download(payload, f"{stem}.pdf", "application/pdf")


def _download(payload: bytes, filename: str, media_type: str) -> StreamingResponse:
    import io
    return StreamingResponse(
        io.BytesIO(payload),
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Filename": filename,
            "Access-Control-Expose-Headers": "Content-Disposition, X-Filename",
        },
    )


async def _analysis_context(body: AnalyzeIn, user) -> Dict[str, Any]:
    """Shared analysis pipeline used by both JSON /report and multi-format /report/{fmt}."""
    text = (body.output or "") + "\n" + body.input
    iocs = extract_iocs(text)
    mitre_hits = mitre_map(text)
    yara = yara_lite_scan(text)
    lolbas = scan_lolbas(text)
    risk = risk_score(mitre_hits, yara, iocs)
    ti_hits = await _lookup_ti_hits(iocs)
    osint = None
    description = None
    verdict = None
    if body.enrich_osint:
        try:
            keys = await load_osint_keys()
            osint = await enrich_iocs(iocs, keys)
        except Exception as e:
            osint = {"error": str(e)}
    if body.describe or body.use_ai_verdict:
        try:
            ai_bundle = await _ai_describe_and_verdict(
                body.input, body.output or "", iocs, mitre_hits, yara, osint or {},
                lolbas=lolbas,
                want_verdict=body.use_ai_verdict, want_describe=body.describe,
            )
            description = ai_bundle.get("description")
            verdict = ai_bundle.get("verdict")
        except Exception as e:
            description = {"error": str(e)}
    # merge AI MITRE with heuristic
    merged_mitre = list(mitre_hits)
    if description and not description.get("error"):
        ai_mitre = description.get("mitre_techniques") or []
        seen_ids = {m["id"] for m in merged_mitre}
        for m in ai_mitre:
            if isinstance(m, dict) and m.get("id") and m["id"] not in seen_ids:
                merged_mitre.append({**m, "source": "ai"})
                seen_ids.add(m["id"])
        for m in merged_mitre:
            m.setdefault("source", "heuristic")
    return {
        "ts": datetime.now(timezone.utc).isoformat(),
        "iocs": iocs, "mitre": merged_mitre, "yara": yara, "lolbas": lolbas,
        "risk": risk, "ti_hits": ti_hits, "osint": osint,
        "description": description, "verdict": verdict,
    }


def _render_csv_report(user, ts, body, ctx) -> str:
    """CSV export — one row per artifact (IOC / MITRE / YARA / LOLBAS / TI-hit)."""
    import csv, io
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["type", "id_or_kind", "value", "severity_or_tactic", "detail", "source"])
    # header metadata
    w.writerow(["META", "generated_at", ts, "", "", "nivxray"])
    w.writerow(["META", "analyst", user["email"], "", "", "nivxray"])
    w.writerow(["META", "verdict", ctx["risk"]["verdict"], f"score={ctx['risk']['score']}", "", "heuristic"])
    fam = (ctx["description"] or {}).get("malware_family") if ctx["description"] else None
    if fam and fam.get("name"):
        w.writerow(["META", "malware_family", fam.get("name", ""), fam.get("confidence", ""), fam.get("rationale", ""), "ai"])
    # MITRE
    for m in ctx["mitre"] or []:
        w.writerow(["MITRE", m.get("id", ""), m.get("technique", ""), m.get("tactic", ""), m.get("evidence", ""), m.get("source", "heuristic")])
    # YARA
    for y in ctx["yara"] or []:
        w.writerow(["YARA", y.get("rule", ""), y.get("match", ""), y.get("severity", ""), y.get("description", ""), "yara-lite"])
    # LOLBAS
    for l in ctx["lolbas"] or []:
        w.writerow(["LOLBAS", l.get("binary", ""), ";".join(l.get("purposes", [])), ";".join(l.get("mitre", [])), l.get("description", ""), l.get("url", "")])
    # IOCs
    for kind, arr in (ctx["iocs"] or {}).items():
        for v in arr or []:
            w.writerow(["IOC", kind, v, "", "", "extracted"])
    # TI hits
    for h in ctx["ti_hits"] or []:
        w.writerow(["TI-HIT", h.get("kind", ""), h.get("value", ""), h.get("severity", ""), ";".join(h.get("tags") or []), h.get("source", "")])
    # OSINT IPs
    if ctx["osint"] and not ctx["osint"].get("error"):
        for ip in ctx["osint"].get("ips") or []:
            geo = ip.get("geo") or {}
            vt = ip.get("virustotal") or {}
            w.writerow(["OSINT-IP", "ip", ip["value"], geo.get("country", ""),
                        f"vt_malicious={vt.get('malicious', 0)};abuseipdb={(ip.get('abuseipdb') or {}).get('abuse_confidence_score', '')};rdns={ip.get('reverse_dns', '')}",
                        ";".join(ctx["osint"].get("sources_used") or [])])
        for d in ctx["osint"].get("domains") or []:
            vt = d.get("virustotal") or {}
            w.writerow(["OSINT-DOMAIN", "domain", d["value"], "", f"vt_malicious={vt.get('malicious', 0)};resolved={','.join(d.get('resolved_ips') or [])}", ""])
        for h in ctx["osint"].get("hashes") or []:
            vt = h.get("virustotal") or {}
            w.writerow(["OSINT-HASH", h["algorithm"], h["value"], "", f"vt_malicious={vt.get('malicious', 0)};label={vt.get('threat_label', '')}", ""])
    return buf.getvalue()


def _render_docx_report(user, ts, body, ctx) -> bytes:
    """Generate a native .docx report (Microsoft Word)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(10)

    # header
    h = doc.add_heading("NivXRay — Decoder & Threat Analysis Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {ts}\n").italic = True
    meta.add_run(f"Analyst: {user['email']}").italic = True

    # verdict
    doc.add_heading("Verdict", 1)
    p = doc.add_paragraph()
    p.add_run(f"{ctx['risk']['verdict']} · heuristic score {ctx['risk']['score']}/100\n").bold = True
    v = ctx["verdict"] or {}
    if v and not v.get("error"):
        p.add_run(f"AI: {v.get('verdict')} · confidence {v.get('confidence')}%\n").bold = True
        p.add_run(v.get("summary", ""))
    fam = (ctx["description"] or {}).get("malware_family") if ctx["description"] else None
    if fam and fam.get("name"):
        pf = doc.add_paragraph()
        r = pf.add_run(f"Malware family: {fam.get('name')} ({fam.get('confidence', '?')} confidence)")
        r.bold = True
        r.font.color.rgb = RGBColor(0xE2, 0x7E, 0x5D)
        doc.add_paragraph(fam.get("rationale", ""))

    # input / decoded
    doc.add_heading("Input (raw)", 1)
    doc.add_paragraph((body.input or "")[:1500])
    doc.add_heading("Decoded output", 1)
    doc.add_paragraph((body.output or "")[:3000])

    # AI describe
    d = ctx["description"] or {}
    if d and not d.get("error"):
        doc.add_heading("AI Analysis", 1)
        if d.get("summary"):
            doc.add_paragraph(d["summary"])
        if d.get("behavior"):
            doc.add_heading("Behavior", 2)
            for b in d["behavior"]:
                doc.add_paragraph(b, style="List Bullet")
        if d.get("ioc_narrative"):
            doc.add_heading("IOC Narrative", 2)
            doc.add_paragraph(d["ioc_narrative"])
        if d.get("attribution_hints"):
            doc.add_heading("Attribution Hints", 2)
            doc.add_paragraph(d["attribution_hints"])
        if d.get("recommended_actions"):
            doc.add_heading("Recommended Actions", 2)
            for a in d["recommended_actions"]:
                doc.add_paragraph(a, style="List Bullet")

    # MITRE table
    if ctx["mitre"]:
        doc.add_heading("MITRE ATT&CK", 1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "ID"; hdr[1].text = "Technique"; hdr[2].text = "Tactic"; hdr[3].text = "Evidence"; hdr[4].text = "Source"
        for m in ctx["mitre"]:
            row = table.add_row().cells
            row[0].text = m.get("id", "")
            row[1].text = m.get("technique", "")
            row[2].text = m.get("tactic", "")
            row[3].text = m.get("evidence", "")
            row[4].text = m.get("source", "heuristic")

    # LOLBAS
    if ctx["lolbas"]:
        doc.add_heading("LOLBAS Matches", 1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Binary"; hdr[1].text = "Purposes"; hdr[2].text = "MITRE"; hdr[3].text = "Description"
        for l in ctx["lolbas"]:
            row = table.add_row().cells
            row[0].text = l.get("binary", "")
            row[1].text = ", ".join(l.get("purposes", []))
            row[2].text = ", ".join(l.get("mitre", []))
            row[3].text = l.get("description", "")

    # YARA
    if ctx["yara"]:
        doc.add_heading("YARA-lite Hits", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Rule"; hdr[1].text = "Severity"; hdr[2].text = "Description"
        for y in ctx["yara"]:
            row = table.add_row().cells
            row[0].text = y.get("rule", "")
            row[1].text = y.get("severity", "")
            row[2].text = y.get("description", "")

    # IOCs
    if ctx["iocs"] and any(v for v in ctx["iocs"].values()):
        doc.add_heading("Extracted IOCs", 1)
        for k, arr in ctx["iocs"].items():
            if not arr: continue
            doc.add_heading(k.upper(), 2)
            for v in arr:
                doc.add_paragraph(v, style="List Bullet")

    # TI hits
    if ctx["ti_hits"]:
        doc.add_heading("Local Threat-Intel Hits", 1)
        for h in ctx["ti_hits"]:
            doc.add_paragraph(f"[{h.get('severity')}] {h.get('kind')} {h.get('value')} — source: {h.get('source')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_pdf_from_html(html: str) -> bytes:
    """Render our HTML report to PDF via xhtml2pdf.
    xhtml2pdf lacks CSS-variable & modern-selector support, so we inline-substitute the token values.
    """
    from xhtml2pdf import pisa
    import re as _re, io
    # substitute CSS custom-property references with concrete hex values
    subs = {
        "var(--bg)": "#101112", "var(--sf)": "#18191b", "var(--inset)": "#0a0a0c",
        "var(--br)": "#2d3135", "var(--ac)": "#4aa890", "var(--warn)": "#e27e5d",
        "var(--hi)": "#d96c6c", "var(--tx)": "#e5e7eb", "var(--dim)": "#8b949e",
    }
    pdf_html = html
    for k, v in subs.items():
        pdf_html = pdf_html.replace(k, v)
    # strip :root { ... } block (contains --var declarations)
    pdf_html = _re.sub(r":root\s*\{[^}]*\}", "", pdf_html)
    # remove flexbox / grid rules xhtml2pdf can't parse (best-effort)
    pdf_html = _re.sub(r"display\s*:\s*flex[^;]*;?", "", pdf_html)
    pdf_html = _re.sub(r"gap\s*:\s*\d+px;?", "", pdf_html)
    buf = io.BytesIO()
    result = pisa.CreatePDF(io.StringIO(pdf_html), dest=buf, encoding="utf-8")
    if result.err:
        raise HTTPException(status_code=500, detail=f"PDF render failed ({result.err} errors)")
    return buf.getvalue()


def _render_text_report(user, ts, body, risk, mitre, yara, lolbas, iocs, ti_hits, osint, description, verdict):
    lines = [
        "NIVXRAY — DECODER & THREAT ANALYSIS REPORT",
        f"Generated: {ts}",
        f"Analyst:   {user['email']}",
        "=" * 68, "",
        f"VERDICT:   {risk['verdict']}   (heuristic score {risk['score']}/100)",
    ]
    if verdict and not verdict.get("error"):
        lines += [f"AI:        {verdict.get('verdict')}   ({verdict.get('confidence')}% confidence)"]
    if description and not description.get("error"):
        fam = description.get("malware_family") or {}
        if fam.get("name"):
            lines += [f"FAMILY:    {fam.get('name')}  ({fam.get('confidence','?')} confidence)"]
    lines += ["", "INPUT (first 400 chars):", (body.input or "")[:400], "",
              "DECODED OUTPUT (first 1500 chars):", (body.output or "")[:1500], ""]
    if description and not description.get("error"):
        lines += ["── AI EXECUTIVE SUMMARY ──", description.get("summary", ""), ""]
        if description.get("behavior"):
            lines += ["── BEHAVIOR ──"] + [f"  · {b}" for b in description["behavior"]] + [""]
        if description.get("ioc_narrative"):
            lines += ["── IOC NARRATIVE ──", description["ioc_narrative"], ""]
        if description.get("attribution_hints"):
            lines += ["── ATTRIBUTION HINTS ──", description["attribution_hints"], ""]
        if description.get("recommended_actions"):
            lines += ["── RECOMMENDED ACTIONS ──"] + [f"  · {a}" for a in description["recommended_actions"]] + [""]
    lines += ["── MITRE ATT&CK ──"]
    for m in mitre or []:
        line = f"  - {m['id']}  {m['technique']}   [{m['tactic']}]"
        if m.get("evidence"): line += f"\n     evidence: {m['evidence']}"
        lines.append(line)
    if not mitre: lines.append("  (none)")
    lines += ["", "── LOLBAS MATCHES ──"]
    for l in lolbas or []:
        lines += [f"  · {l['binary']}   purposes={','.join(l['purposes'])}   mitre={','.join(l['mitre'])}",
                  f"     {l['description']}",
                  f"     snippet: {l['snippet']}"]
    if not lolbas: lines.append("  (none)")
    lines += ["", "── YARA-LITE HITS ──"]
    for y in yara or []:
        lines.append(f"  - [{y['severity'].upper()}] {y['rule']}: {y['description']}")
    if not yara: lines.append("  (none)")
    lines += ["", "── IOCs ──"]
    for k, v in (iocs or {}).items():
        if v:
            lines.append(f"  {k}:")
            for item in v: lines.append(f"    - {item}")
    lines += ["", "── LOCAL THREAT-INTEL HITS ──"]
    for h in ti_hits or []:
        lines.append(f"  - [{h['severity']}] {h['kind']} {h['value']}  (source: {h['source']})")
    if not ti_hits: lines.append("  (none)")
    if osint:
        lines += ["", "── OSINT ENRICHMENT ──", f"  sources: {', '.join(osint.get('sources_used') or [])}"]
        for ip in osint.get("ips", []) or []:
            geo = ip.get("geo") or {}
            lines.append(f"  IP {ip['value']}: {geo.get('country','?')} / {geo.get('isp','?')}"
                         + (f"  rDNS={ip.get('reverse_dns')}" if ip.get('reverse_dns') else ""))
            if ip.get("virustotal"):
                lines.append(f"     VT: {ip['virustotal'].get('malicious',0)} malicious")
            if ip.get("abuseipdb"):
                lines.append(f"     AbuseIPDB: {ip['abuseipdb'].get('abuse_confidence_score',0)}% confidence")
    lines += ["", "=" * 68, "End of report."]
    return "\n".join(lines)


def _render_html_report(user, ts, body, risk, mitre, yara, lolbas, iocs, ti_hits, osint, description, verdict):
    from html import escape as _e
    def block(title, body_html):
        return f'<section><h2>{_e(title)}</h2><div class="card">{body_html}</div></section>'
    def _sev_class(s): return f'sev-{_e(s)}'
    fam = (description or {}).get("malware_family") or {}
    parts = []
    parts.append(f'''<!doctype html><html><head><meta charset="utf-8"><title>NivXRay report</title>
<style>
:root {{ --bg:#101112; --sf:#18191b; --inset:#0a0a0c; --br:#2d3135; --ac:#4aa890; --warn:#e27e5d; --hi:#d96c6c; --tx:#e5e7eb; --dim:#8b949e; }}
* {{ box-sizing: border-box; }}
body {{ background:var(--bg); color:var(--tx); font-family: Chivo, ui-sans-serif, sans-serif; margin:0; padding:32px; line-height:1.55; }}
.mono {{ font-family: 'JetBrains Mono', ui-monospace, monospace; }}
.hdr {{ display:flex; align-items:center; gap:14px; padding-bottom:16px; border-bottom:1px solid var(--br); }}
.hdr .logo {{ width:24px; height:24px; border:1px solid var(--ac); position:relative; }}
.hdr .logo::before {{ content:''; position:absolute; inset:6px; background:var(--ac); }}
.hdr h1 {{ font-weight:900; letter-spacing:0.14em; margin:0; }}
.hdr h1 span {{ color:var(--ac); }}
.meta {{ color:var(--dim); font-size:12px; margin-left:auto; text-align:right; }}
section {{ margin-top:26px; }}
section h2 {{ color:var(--ac); font-size:11px; letter-spacing:0.22em; margin:0 0 10px 0; }}
.card {{ background:var(--sf); border:1px solid var(--br); padding:16px; }}
.badge {{ display:inline-block; padding:3px 7px; border:1px solid var(--br); font-family:'JetBrains Mono',monospace; font-size:10px; letter-spacing:0.06em; margin-right:6px; }}
.badge.hi {{ color:var(--hi); border-color:var(--hi); background:rgba(217,108,108,0.1); }}
.badge.med {{ color:var(--warn); border-color:var(--warn); background:rgba(226,126,93,0.1); }}
.badge.low {{ color:#c0ca33; border-color:#c0ca33; background:rgba(192,202,51,0.1); }}
.badge.safe {{ color:var(--ac); border-color:var(--ac); background:rgba(74,168,144,0.1); }}
.badge.critical {{ color:var(--hi); border-color:var(--hi); background:rgba(217,108,108,0.2); }}
.badge.high {{ color:var(--hi); border-color:var(--hi); }}
.badge.medium {{ color:var(--warn); border-color:var(--warn); }}
pre {{ background:var(--inset); padding:12px; border:1px solid var(--br); overflow-x:auto; font-size:11px; color:var(--tx); white-space:pre-wrap; word-break:break-all; margin:0; }}
table {{ width:100%; border-collapse:collapse; font-family:'JetBrains Mono',monospace; font-size:11px; }}
th,td {{ text-align:left; padding:6px 8px; border-bottom:1px solid var(--br); vertical-align:top; }}
th {{ color:var(--dim); font-weight:700; letter-spacing:0.12em; font-size:10px; }}
ul {{ margin:0; padding-left:20px; }}
ul li {{ margin:4px 0; }}
.verdict {{ display:flex; align-items:center; gap:14px; }}
.verdict .score {{ font-size:40px; font-weight:900; color:var(--ac); }}
.warn {{ color:var(--warn); }}
.hi {{ color:var(--hi); }}
</style></head><body>
<div class="hdr">
  <div class="logo"></div>
  <h1>NIVX<span>RAY</span> · DECODER & THREAT ANALYSIS REPORT</h1>
  <div class="meta">Generated: {_e(ts)}<br>Analyst: {_e(user['email'])}</div>
</div>''')
    # Verdict
    vv = verdict or {}
    fam_line = ""
    if fam.get("name"):
        fam_line = f"<div class='mono' style='margin-top:6px;'>Family: <span class='warn'>{_e(fam.get('name'))}</span> <span class='badge {fam.get('confidence','low')}'>{_e(fam.get('confidence','?'))} confidence</span></div>"
    parts.append(block("VERDICT", f"""
<div class='verdict'>
  <div class='score'>{_e(str(risk['score']))}<span style='font-size:16px; color:var(--dim);'>/100</span></div>
  <div>
    <div><span class='badge {_e(risk['level'])}'>{_e(risk['verdict'])}</span>
    {"<span class='badge hi'>AI: " + _e(vv.get('verdict','?')) + " " + _e(str(vv.get('confidence','?'))) + "%</span>" if vv and not vv.get('error') else ''}</div>
    {fam_line}
    {"<div class='mono' style='margin-top:6px; color:var(--dim);'>" + _e(vv.get('summary','')) + "</div>" if vv and not vv.get('error') else ''}
  </div>
</div>"""))

    # Input & decoded output
    parts.append(block("INPUT", f"<pre class='mono'>{_e((body.input or '')[:2000])}</pre>"))
    parts.append(block("DECODED OUTPUT", f"<pre class='mono'>{_e((body.output or '')[:4000])}</pre>"))

    # AI describe
    if description and not description.get("error"):
        d = description
        blocks = []
        if d.get("summary"): blocks.append(f"<p class='mono'>{_e(d['summary'])}</p>")
        if d.get("behavior"): blocks.append("<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>BEHAVIOR</h3><ul class='mono'>" + "".join(f"<li>{_e(b)}</li>" for b in d['behavior']) + "</ul>")
        if d.get("ioc_narrative"): blocks.append(f"<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>IOC NARRATIVE</h3><p class='mono' style='color:var(--dim);'>{_e(d['ioc_narrative'])}</p>")
        if d.get("attribution_hints"): blocks.append(f"<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>ATTRIBUTION HINTS</h3><p class='mono' style='color:var(--dim);'>{_e(d['attribution_hints'])}</p>")
        if d.get("recommended_actions"): blocks.append("<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>RECOMMENDED ACTIONS</h3><ul class='mono' style='color:var(--ac);'>" + "".join(f"<li>{_e(a)}</li>" for a in d['recommended_actions']) + "</ul>")
        parts.append(block("AI ANALYSIS", "".join(blocks)))

    # MITRE table
    if mitre:
        rows = "".join(
            f"<tr><td><a href='https://attack.mitre.org/techniques/{m['id'].replace('.','/')}/' target='_blank' style='color:var(--ac);'>{_e(m['id'])}</a></td><td>{_e(m.get('technique',''))}</td><td>{_e(m.get('tactic',''))}</td><td>{_e(m.get('evidence',''))}</td><td><span class='badge'>{_e(m.get('source','heuristic'))}</span></td></tr>"
            for m in mitre
        )
        parts.append(block("MITRE ATT&CK", f"<table><tr><th>ID</th><th>Technique</th><th>Tactic</th><th>Evidence</th><th>Source</th></tr>{rows}</table>"))

    # LOLBAS
    if lolbas:
        rows = "".join(
            f"<tr><td class='warn'>{_e(l['binary'])}</td><td>{', '.join(_e(p) for p in l['purposes'])}</td><td>{', '.join(_e(t) for t in l['mitre'])}</td><td>{_e(l['description'])}</td><td><a style='color:var(--ac);' href='{_e(l['url'])}' target='_blank'>docs</a></td></tr>"
            for l in lolbas
        )
        parts.append(block("LOLBAS", f"<table><tr><th>Binary</th><th>Purposes</th><th>MITRE</th><th>Description</th><th></th></tr>{rows}</table>"))

    # YARA
    if yara:
        rows = "".join(
            f"<tr><td>{_e(y['rule'])}</td><td><span class='badge {_e(y['severity'])}'>{_e(y['severity'])}</span></td><td>{_e(y['description'])}</td><td class='mono' style='color:var(--dim);'>{_e(y['match'][:80])}</td></tr>"
            for y in yara
        )
        parts.append(block("YARA-LITE HITS", f"<table><tr><th>Rule</th><th>Severity</th><th>Description</th><th>Match</th></tr>{rows}</table>"))

    # IOCs
    ioc_rows = []
    for k, v in (iocs or {}).items():
        for item in v or []:
            ioc_rows.append(f"<tr><td>{_e(k)}</td><td class='mono'>{_e(item)}</td></tr>")
    if ioc_rows:
        parts.append(block("EXTRACTED IOCs", f"<table><tr><th>Kind</th><th>Value</th></tr>{''.join(ioc_rows)}</table>"))

    # TI hits
    if ti_hits:
        rows = "".join(
            f"<tr><td>{_e(h.get('kind',''))}</td><td class='mono'>{_e(h.get('value',''))}</td><td><span class='badge {_e(h.get('severity','low'))}'>{_e(h.get('severity',''))}</span></td><td>{_e(h.get('source',''))}</td></tr>"
            for h in ti_hits
        )
        parts.append(block("LOCAL THREAT-INTEL HITS", f"<table><tr><th>Kind</th><th>Value</th><th>Severity</th><th>Source</th></tr>{rows}</table>"))

    # OSINT
    if osint and not osint.get("error"):
        html_bits = [f"<div class='mono' style='color:var(--dim);margin-bottom:8px;'>Sources: {_e(', '.join(osint.get('sources_used') or []))}</div>"]
        for ip in osint.get("ips") or []:
            geo = ip.get("geo") or {}
            vt = ip.get("virustotal") or {}
            ab = ip.get("abuseipdb") or {}
            html_bits.append(f"<div class='card' style='margin-bottom:8px;'><b class='mono' style='color:var(--ac);'>{_e(ip['value'])}</b> — {_e(geo.get('country',''))} · {_e(geo.get('isp',''))}"
                             + (f" · rDNS={_e(ip.get('reverse_dns',''))}" if ip.get('reverse_dns') else "")
                             + (f"<br>VT: <span class='hi'>{vt.get('malicious',0)} malicious</span>, {vt.get('suspicious',0)} suspicious" if vt else "")
                             + (f"<br>AbuseIPDB: {ab.get('abuse_confidence_score',0)}% confidence, {ab.get('total_reports',0)} reports" if ab else "")
                             + "</div>")
        parts.append(block("OSINT ENRICHMENT", "".join(html_bits)))

    parts.append('</body></html>')
    return "".join(parts)


# =============================================================================
# Admin — Settings (OSINT API Keys) + Users listing
# =============================================================================
@api.get("/admin/osint/services")
async def get_osint_services(user=Depends(require_admin)):
    keys = await load_osint_keys()
    return [
        {**s, "configured": bool(keys.get(s["id"])), "masked_key": _mask(keys.get(s["id"], ""))}
        for s in OSINT_SERVICES
    ]


@api.put("/admin/osint/settings")
async def update_osint_settings(body: SettingsUpdateIn, user=Depends(require_admin)):
    existing = await load_osint_keys()
    merged = {**existing}
    for svc_id, key in body.keys.items():
        if svc_id not in [s["id"] for s in OSINT_SERVICES]:
            continue
        # empty string => remove; otherwise set
        if key == "":
            merged.pop(svc_id, None)
        else:
            merged[svc_id] = key.strip()
    await db.settings.update_one(
        {"_id": "osint_keys"},
        {"$set": {"keys": merged, "updated_by": user["email"], "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"ok": True, "configured_services": [k for k, v in merged.items() if v]}


@api.post("/admin/osint/test/{service_id}")
async def test_osint(service_id: str, user=Depends(require_admin)):
    keys = await load_osint_keys()
    key = keys.get(service_id)
    if not key:
        raise HTTPException(status_code=400, detail=f"No API key configured for {service_id}")
    # dispatch a minimal test
    import httpx
    async with httpx.AsyncClient(timeout=8.0, headers={"User-Agent": "NivXRay/1.0"}) as c:
        try:
            if service_id == "virustotal":
                r = await c.get("https://www.virustotal.com/api/v3/ip_addresses/8.8.8.8", headers={"x-apikey": key})
            elif service_id == "abuseipdb":
                r = await c.get("https://api.abuseipdb.com/api/v2/check",
                                headers={"Key": key, "Accept": "application/json"},
                                params={"ipAddress": "8.8.8.8", "maxAgeInDays": 30})
            elif service_id == "shodan":
                r = await c.get("https://api.shodan.io/api-info", params={"key": key})
            elif service_id == "greynoise":
                r = await c.get("https://api.greynoise.io/v3/community/8.8.8.8", headers={"key": key})
            elif service_id == "urlscan":
                r = await c.get("https://urlscan.io/user/quotas/", headers={"API-Key": key})
            elif service_id == "otx":
                r = await c.get("https://otx.alienvault.com/api/v1/user/me", headers={"X-OTX-API-KEY": key})
            elif service_id == "ipinfo":
                r = await c.get("https://ipinfo.io/8.8.8.8", params={"token": key})
            elif service_id == "hybrid_analysis":
                r = await c.get("https://www.hybrid-analysis.com/api/v2/key/current",
                                headers={"api-key": key, "user-agent": "Falcon Sandbox"})
            else:
                raise HTTPException(status_code=400, detail="Unknown service")
            return {"ok": r.status_code < 400, "status_code": r.status_code, "body_snippet": r.text[:200]}
        except HTTPException:
            raise
        except Exception as e:
            return {"ok": False, "error": str(e)}


@api.get("/admin/users")
async def list_users(user=Depends(require_admin)):
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(50)
    return users


@api.get("/admin/stats")
async def admin_stats(user=Depends(require_admin)):
    total_shares = await db.shares.count_documents({})
    total_users = await db.users.count_documents({})
    total_iocs = await db.iocs.count_documents({})
    keys = await load_osint_keys()
    return {
        "total_shares": total_shares,
        "total_users": total_users,
        "total_iocs": total_iocs,
        "configured_osint_services": len([v for v in keys.values() if v]),
        "total_operations": len(OPERATIONS),
        "lolbas": lolbas_status(),
    }


# =============================================================================
# LOLBAS catalog admin — sync + status
# =============================================================================
# =============================================================================
# Model Studio — custom detection rules, decode recipes, AI personas, providers
# =============================================================================
class ModelIn(BaseModel):
    kind: str
    name: str
    enabled: bool = True
    config: Dict[str, Any] = Field(default_factory=dict)


class ModelPatchIn(BaseModel):
    name: Optional[str] = None
    enabled: Optional[bool] = None
    config: Optional[Dict[str, Any]] = None


class ModelTestIn(BaseModel):
    sample: str


@api.get("/admin/models")
async def list_admin_models(kind: Optional[str] = None, user=Depends(require_admin)):
    if kind and kind not in ms.MODEL_KINDS:
        raise HTTPException(status_code=400, detail=f"invalid kind: {kind}")
    return await ms.list_models(db, kind=kind)


@api.get("/admin/models/catalog")
async def model_catalog(user=Depends(require_admin)):
    """Static catalog needed by the Model Studio UI (available ops + kinds)."""
    return {
        "kinds": list(ms.MODEL_KINDS),
        "operations": sorted(OPERATIONS.keys()),
        "providers": [
            {"provider": "anthropic", "model": "claude-sonnet-4-5-20250929", "label": "Claude Sonnet 4.5"},
            {"provider": "openai", "model": "gpt-5.2", "label": "GPT-5.2"},
            {"provider": "google", "model": "gemini-3-pro", "label": "Gemini 3 Pro"},
        ],
    }


@api.post("/admin/models")
async def create_admin_model(body: ModelIn, user=Depends(require_admin)):
    try:
        return await ms.create_model(
            db, kind=body.kind, name=body.name, config=body.config or {},
            created_by=user.get("email"), enabled=body.enabled,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api.put("/admin/models/{model_id}")
async def update_admin_model(model_id: str, body: ModelPatchIn, user=Depends(require_admin)):
    try:
        patch = {k: v for k, v in body.model_dump().items() if v is not None}
        updated = await ms.update_model(db, model_id, patch)
        if not updated:
            raise HTTPException(status_code=404, detail="model not found")
        return updated
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api.delete("/admin/models/{model_id}")
async def delete_admin_model(model_id: str, user=Depends(require_admin)):
    try:
        ok = await ms.delete_model(db, model_id)
        if not ok:
            raise HTTPException(status_code=404, detail="model not found")
        return {"deleted": True}
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))


@api.post("/admin/models/{model_id}/test")
async def test_admin_model(model_id: str, body: ModelTestIn, user=Depends(require_admin)):
    """Test a model against a sample input. Returns a kind-specific result payload."""
    m = await ms.get_model(db, model_id)
    if not m:
        raise HTTPException(status_code=404, detail="model not found")
    kind = m["kind"]
    cfg = m.get("config") or {}
    sample = body.sample or ""
    if kind == "detection_rule":
        hits = ms.scan_custom_rules(sample, [{
            "binary_regex": cfg.get("binary_regex", ""),
            "argv": cfg.get("argv_regex"),
            "purposes": cfg.get("purposes") or ["Custom"],
            "mitre": cfg.get("mitre") or [],
            "desc": cfg.get("description") or m["name"],
            "url": "", "source": f"custom:{m['id']}",
            "name": m["name"], "severity": cfg.get("severity", "medium"),
            "model_id": m["id"],
        }])
        return {"kind": kind, "matched": bool(hits), "hits": hits}
    if kind == "decode_recipe":
        try:
            matched = bool(re.search(cfg.get("match_regex", ""), sample, re.IGNORECASE | re.DOTALL))
        except re.error as e:
            raise HTTPException(status_code=400, detail=f"invalid match_regex: {e}")
        steps_out: List[Dict[str, Any]] = []
        output = sample
        if matched:
            for step in (cfg.get("ops") or []):
                op = step.get("op")
                if op not in OPERATIONS:
                    steps_out.append({"op": op, "error": "unknown op"})
                    break
                try:
                    output = run_operation(op, output, step.get("args") or {})
                    steps_out.append({"op": op, "output_preview": (output or "")[:300]})
                except Exception as e:
                    steps_out.append({"op": op, "error": str(e)})
                    break
        return {"kind": kind, "matched": matched, "steps": steps_out, "output": output}
    if kind == "ai_persona":
        return {"kind": kind,
                "system_prompt_preview": (cfg.get("system_prompt") or "")[:800],
                "sample_would_be_sent_as": (sample or "(no sample provided)")[:400]}
    if kind == "ai_provider":
        return {"kind": kind, "provider": cfg.get("provider"), "model": cfg.get("model"),
                "note": "Provider connectivity is verified at Auto-Investigate time via the Emergent Universal LLM Key."}
    if kind == "playbook":
        return {"kind": kind,
                "applies_to": cfg.get("applies_to") or ["ai"],
                "body_preview": (cfg.get("body") or "")[:800],
                "note": "This playbook is auto-appended to every AI investigation. Trigger AUTO-INVESTIGATE on a sample to see the effect."}
    if kind == "training_note":
        return {"kind": kind,
                "body_preview": (cfg.get("body") or "")[:800],
                "note": "This training note is ALWAYS PREPENDED to every AI investigation system prompt (ranked above playbooks). Analyst 👍/👎 votes adjust its ordering for future prompts."}
    raise HTTPException(status_code=400, detail=f"unknown kind: {kind}")


# =============================================================================
# Malware Sample Library — regression + benchmark + coverage dashboard
# =============================================================================
class SampleIn(BaseModel):
    name: str
    raw_input: str
    expected_output: str
    categories: Optional[List[str]] = None
    tags: Optional[List[str]] = None
    expected_mitre: Optional[List[str]] = None
    expected_iocs: Optional[List[str]] = None
    difficulty: Optional[str] = "medium"
    source_url: Optional[str] = None
    notes: Optional[str] = None


class SamplePatchIn(BaseModel):
    name: Optional[str] = None
    raw_input: Optional[str] = None
    expected_output: Optional[str] = None
    categories: Optional[List[str]] = None
    tags: Optional[List[str]] = None
    expected_mitre: Optional[List[str]] = None
    expected_iocs: Optional[List[str]] = None
    difficulty: Optional[str] = None
    source_url: Optional[str] = None
    notes: Optional[str] = None


class SampleBulkIn(BaseModel):
    samples: List[SampleIn]


@api.get("/admin/samples")
async def list_samples_endpoint(category: Optional[str] = None, user=Depends(require_admin)):
    return await sl.list_samples(db, category=category)


@api.get("/admin/samples/dashboard")
async def samples_dashboard(user=Depends(require_admin)):
    return await sl.dashboard_snapshot(db)


@api.get("/admin/samples/{sid}")
async def get_sample_endpoint(sid: str, user=Depends(require_admin)):
    s = await sl.get_sample(db, sid)
    if not s:
        raise HTTPException(status_code=404, detail="sample not found")
    return s


@api.post("/admin/samples")
async def create_sample_endpoint(body: SampleIn, user=Depends(require_admin)):
    try:
        return await sl.create_sample(db, body.model_dump())
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api.post("/admin/samples/bulk")
async def bulk_create_samples(body: SampleBulkIn, user=Depends(require_admin)):
    created, failed = [], []
    for s in body.samples:
        try:
            created.append(await sl.create_sample(db, s.model_dump()))
        except Exception as e:
            failed.append({"name": s.name, "error": str(e)})
    return {"created": len(created), "failed": len(failed), "items": created, "errors": failed}


@api.put("/admin/samples/{sid}")
async def update_sample_endpoint(sid: str, body: SamplePatchIn, user=Depends(require_admin)):
    patch = {k: v for k, v in body.model_dump().items() if v is not None}
    updated = await sl.update_sample(db, sid, patch)
    if not updated:
        raise HTTPException(status_code=404, detail="sample not found")
    return updated


@api.delete("/admin/samples/{sid}")
async def delete_sample_endpoint(sid: str, user=Depends(require_admin)):
    try:
        ok = await sl.delete_sample(db, sid)
        if not ok:
            raise HTTPException(status_code=404, detail="sample not found")
        return {"deleted": True}
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))


@api.post("/admin/samples/{sid}/benchmark")
async def benchmark_one_endpoint(sid: str, user=Depends(require_admin)):
    try:
        return await sl.benchmark_one(db, sid, smart_decode, magic_decode)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@api.post("/admin/samples/benchmark/all")
async def benchmark_all_endpoint(user=Depends(require_admin)):
    return await sl.benchmark_all(db, smart_decode, magic_decode)


async def _nightly_benchmark_loop():
    """Runs the full sample-library benchmark every 24h in the background.

    Failures are logged but never crash the loop; the next tick just tries again.
    """
    # small startup delay so first startup isn't burdened
    await asyncio.sleep(300)
    while True:
        try:
            r = await sl.benchmark_all(db, smart_decode, magic_decode)
            log.info("nightly benchmark: %d samples · %d passed (%.1f%%)",
                     r.get("total", 0), r.get("passed", 0), r.get("pass_pct", 0.0))
        except Exception as e:
            log.warning("nightly benchmark failed: %s", e)
        await asyncio.sleep(24 * 60 * 60)


@api.get("/admin/lolbas/status")
async def get_lolbas_status(user=Depends(require_admin)):
    return lolbas_status()


@api.post("/admin/lolbas/sync")
async def sync_lolbas_catalog(user=Depends(require_admin)):
    """Force-refresh the LOLBAS catalog from lolbas-project.github.io."""
    return await lolbas_refresh(db)


# =============================================================================
# Threat Intelligence — IOC Database & bulk-feed sync
# =============================================================================
async def _iocs_indexes():
    """Ensure indexes on the iocs collection."""
    try:
        await db.iocs.create_index([("kind", 1), ("value", 1), ("source", 1)], unique=True, name="uniq_ioc")
        await db.iocs.create_index([("source", 1)])
        await db.iocs.create_index([("severity", 1)])
        await db.iocs.create_index([("last_seen", -1)])
    except Exception:
        pass


@api.get("/threat-intel/sources")
async def ti_sources(user=Depends(get_current_user)):
    keys = await load_osint_keys()
    # per-source metadata (last sync, count)
    meta_docs = {m["_id"]: m async for m in db.ti_source_meta.find({})}
    out = []
    for s in FEED_SOURCES:
        needs = s.get("needs_key")
        configured = (needs is None) or bool(keys.get(needs))
        m = meta_docs.get(s["id"]) or {}
        out.append({
            **s,
            "configured": configured,
            "last_sync": m.get("last_sync"),
            "last_status": m.get("last_status"),
            "last_new": m.get("last_new", 0),
            "last_updated": m.get("last_updated", 0),
            "last_error": m.get("last_error"),
            "total_indicators": m.get("total_indicators", 0),
        })
    return out


@api.get("/threat-intel/stats")
async def ti_stats(user=Depends(get_current_user)):
    total = await db.iocs.count_documents({})
    critical = await db.iocs.count_documents({"severity": "critical"})
    high = await db.iocs.count_documents({"severity": "high"})
    medium = await db.iocs.count_documents({"severity": "medium"})
    low = await db.iocs.count_documents({"severity": "low"})
    by_kind = {}
    for k in ("ip", "domain", "url", "md5", "sha1", "sha256"):
        by_kind[k] = await db.iocs.count_documents({"kind": k})
    return {"total": total, "critical": critical, "high": high, "medium": medium, "low": low, "by_kind": by_kind}


async def _apply_iocs(iocs: List[Dict[str, Any]], source_id: str) -> Dict[str, int]:
    """Upsert a batch of IOC docs. Returns {'new': int, 'updated': int}."""
    new_count = 0
    upd_count = 0
    for doc in iocs:
        key = {"kind": doc["kind"], "value": doc["value"], "source": source_id}
        existing = await db.iocs.find_one(key, {"_id": 1})
        update = {
            "$set": {"severity": doc["severity"], "tags": doc["tags"], "extra": doc["extra"], "last_seen": doc["last_seen"]},
            "$setOnInsert": {"first_seen": doc["first_seen"]},
        }
        r = await db.iocs.update_one(key, update, upsert=True)
        if r.upserted_id is not None or existing is None:
            new_count += 1
        else:
            upd_count += 1
    return {"new": new_count, "updated": upd_count}


@api.post("/threat-intel/sync/{source_id}")
async def ti_sync_one(source_id: str, user=Depends(require_admin)):
    src = next((s for s in FEED_SOURCES if s["id"] == source_id), None)
    if not src:
        raise HTTPException(status_code=404, detail="Unknown source")
    if not src.get("bulk"):
        raise HTTPException(status_code=400, detail="This source is lookup-only (no bulk feed)")
    keys = await load_osint_keys()
    result = await sync_source(source_id, keys)
    ts = datetime.now(timezone.utc).isoformat()
    if result.get("error"):
        await db.ti_source_meta.update_one(
            {"_id": source_id},
            {"$set": {"last_sync": ts, "last_status": "error", "last_error": result["error"]}},
            upsert=True,
        )
        return {"ok": False, "error": result["error"], "source": source_id}
    counts = await _apply_iocs(result["iocs"], source_id)
    total = await db.iocs.count_documents({"source": source_id})
    await db.ti_source_meta.update_one(
        {"_id": source_id},
        {"$set": {"last_sync": ts, "last_status": "ok", "last_error": None,
                  "last_new": counts["new"], "last_updated": counts["updated"],
                  "total_indicators": total}},
        upsert=True,
    )
    return {"ok": True, "source": source_id, "fetched": len(result["iocs"]), **counts, "total_indicators": total}


@api.post("/threat-intel/sync-all")
async def ti_sync_all(user=Depends(require_admin)):
    keys = await load_osint_keys()
    bulk_sources = [s for s in FEED_SOURCES if s.get("bulk")]
    async def _one(src):
        return src["id"], await sync_source(src["id"], keys)
    results = await asyncio.gather(*[_one(s) for s in bulk_sources], return_exceptions=True)
    summary = []
    ts = datetime.now(timezone.utc).isoformat()
    for r in results:
        if isinstance(r, Exception):
            summary.append({"source": "?", "error": str(r), "ok": False})
            continue
        sid, res = r
        if res.get("error"):
            await db.ti_source_meta.update_one(
                {"_id": sid},
                {"$set": {"last_sync": ts, "last_status": "error", "last_error": res["error"]}},
                upsert=True,
            )
            summary.append({"source": sid, "ok": False, "error": res["error"]})
            continue
        counts = await _apply_iocs(res["iocs"], sid)
        total = await db.iocs.count_documents({"source": sid})
        await db.ti_source_meta.update_one(
            {"_id": sid},
            {"$set": {"last_sync": ts, "last_status": "ok", "last_error": None,
                      "last_new": counts["new"], "last_updated": counts["updated"],
                      "total_indicators": total}},
            upsert=True,
        )
        summary.append({"source": sid, "ok": True, "fetched": len(res["iocs"]), **counts, "total_indicators": total})
    return {"results": summary, "ts": ts}


@api.get("/threat-intel/iocs")
async def ti_iocs(user=Depends(get_current_user), q: str = "", kind: str = "", source: str = "", severity: str = "", limit: int = 100, skip: int = 0):
    query: Dict[str, Any] = {}
    if kind: query["kind"] = kind
    if source: query["source"] = source
    if severity: query["severity"] = severity
    if q:
        query["value"] = {"$regex": re.escape(q), "$options": "i"}
    cur = db.iocs.find(query, {"_id": 0}).sort("last_seen", -1).skip(max(0, skip)).limit(max(1, min(500, limit)))
    docs = await cur.to_list(limit)
    total = await db.iocs.count_documents(query)
    return {"total": total, "items": docs}


@api.get("/threat-intel/lookup/{value}")
async def ti_lookup(value: str, user=Depends(get_current_user)):
    """Return every stored IOC that matches this exact value (across all sources)."""
    docs = await db.iocs.find({"value": value}, {"_id": 0}).to_list(50)
    return {"value": value, "hits": docs}


# =============================================================================
# Load Example Presets
# =============================================================================
EXAMPLES = [
    {
        "id": "powershell-encoded",
        "label": "PowerShell -EncodedCommand",
        "input": "powershell.exe -NoP -NonI -W Hidden -Enc SQBFAFgAKABOAGUAdwAtAE8AYgBqAGUAYwB0ACAATgBlAHQALgBXAGUAYgBDAGwAaQBlAG4AdAApAC4ARABvAHcAbgBsAG8AYQBkAFMAdAByAGkAbgBnACgAJwBoAHQAdABwADoALwAvADEAOQAyAC4AMQA2ADgALgAxAC4AMQAvAHAALgBwAHMAMQAnACkA",
    },
    {
        "id": "ransomware-note",
        "label": "Ransomware Note",
        "input": "!!! YOUR FILES HAVE BEEN ENCRYPTED !!!\nAll your important documents, photos, databases and other files have been encrypted with military-grade AES-256.\n\nTo restore your files you must pay 0.75 BTC to the following address within 72 hours:\n\nBTC ADDRESS: bc1qxy2kgdygjrsqtzq2n0yrf2493p83kkfjhx0wlh\n\nContact us via Tor: http://ransomxyz1abcdef23456789ghijklmn.onion\nEmail: recover-your-files@protonmail.com\n\nDo NOT rename encrypted files. Do NOT try to decrypt with third-party software.\n",
    },
    {
        "id": "defanged-iocs",
        "label": "Defanged IOCs Bundle",
        "input": "IOC dump from IR ticket #4421:\n\nURLs:\n  hxxps://malicious-cdn[.]example[.]com/payload[.]exe\n  hxxp://phish[.]login-microsoft-secure[.]net/auth\n\nIPs:\n  185[.]220[.]101[.]45\n  45[.]137[.]21[.]9\n\nEmails:\n  attacker[@]evilcorp[.]ru\n  admin[@]phish[.]login-microsoft-secure[.]net\n\nHashes:\n  MD5:    e10adc3949ba59abbe56e057f20f883e\n  SHA256: 5e884898da28047151d0e56f8dc6292773603d0d6aabbdd62a11ef721d1542d8\n",
    },
    {
        "id": "nested-base64-gzip",
        "label": "Nested Base64 → gzip",
        "input": "H4sIAIQ5VWoC/xXKyRWAIAwFwFZ+A3ryWYkNBBIRF4LEvXr1PNMNgnWPfoIreib0emHcl2zQQwq2j2d6brCGGt0QDZnuWYlxkiE8MVdel1zETPjvCY5M2qaS5JWF6xdjITdRYgAAAA==",
    },
    {
        "id": "url-encoded-xss",
        "label": "URL-encoded XSS",
        "input": "%3Cscript%3Ealert(String.fromCharCode(88%2C83%2C83))%3C%2Fscript%3E",
    },
]


# =============================================================================
# App wiring
# =============================================================================
app.include_router(api)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def _startup():
    await seed_admin()
    await _iocs_indexes()
    # LOLBAS: load persisted cache, then trigger a background refresh if stale (>7d)
    await lolbas_load(db)
    asyncio.create_task(lolbas_maybe_refresh(db))
    # Model Studio: indexes + seed built-in personas/providers/examples
    await ms.ensure_indexes(db)
    await ms.seed_builtins(db)
    await ms.ensure_vote_indexes(db)
    # Sample Library: indexes + seed 15 built-in samples + start nightly benchmark
    await sl.ensure_indexes(db)
    await sl.seed_builtins(db)
    asyncio.create_task(_nightly_benchmark_loop())


@app.on_event("shutdown")
async def _shutdown():
    client.close()
