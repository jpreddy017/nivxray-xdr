"""Documents / Case Vault router (Feb 2026).

In-app file storage so analysts can upload artefacts of any supported type
directly from the browser when the chat-side upload is unavailable.

Storage: MongoDB GridFS (`documents` bucket) — no extra infra.
Auth: user-scoped, admin can list all via `?all=true`.

Endpoints:
    POST  /api/documents/upload               multipart file(s) upload
    GET   /api/documents                        list current user's uploads
    GET   /api/documents/{id}                   metadata for one doc
    GET   /api/documents/{id}/download          raw bytes stream
    GET   /api/documents/{id}/preview           safe text preview (best-effort)
    DELETE /api/documents/{id}
    POST  /api/documents/{id}/ingest-fixture    for JSON — feed into ir_export_to_fixture
    POST  /api/documents/{id}/re-investigate    for text/JSON — pipe through /decode/smart
"""
from __future__ import annotations

import io
import json
import mimetypes
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from bson import ObjectId
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, Body
from fastapi.responses import StreamingResponse
from motor.motor_asyncio import AsyncIOMotorGridFSBucket
from pydantic import BaseModel

from deps import db, client, DB_NAME, get_current_user

router = APIRouter()

# ── Configuration ────────────────────────────────────────────────────────────
BUCKET_NAME = "documents"
MAX_UPLOAD_BYTES = 25 * 1024 * 1024  # 25 MB per file
ALLOWED_EXTENSIONS = {
    "pdf", "doc", "docx", "csv", "xls", "xlsx",
    "json", "jsonl", "txt", "log", "eml", "msg",
    "html", "htm", "xml", "md", "yml", "yaml",
    "ps1", "bat", "sh", "py", "js", "vbs",
}
PREVIEWABLE_TEXT_EXTS = {
    "json", "jsonl", "txt", "log", "eml", "html", "htm", "xml",
    "md", "yml", "yaml", "csv", "ps1", "bat", "sh", "py", "js", "vbs",
}


# ── Feb 2026 · dict-safe accessor for GridFSFile OR .files dict ────
def _gf(f, key, default=None):
    """Read a field from a motor GridFSFile object OR a raw `.files` dict.
    Motor's `open_download_stream` returns GridFSFile, but `find(...)` returns
    dicts, which broke the docs upload flow (AttributeError: metadata)."""
    if isinstance(f, dict):
        # Filename lives under "filename" in both; upload_date is `uploadDate`
        alt = {"upload_date": "uploadDate", "_id": "_id"}
        return f.get(alt.get(key, key), default)
    return getattr(f, key, default)


def _bucket() -> AsyncIOMotorGridFSBucket:
    return AsyncIOMotorGridFSBucket(client[DB_NAME], bucket_name=BUCKET_NAME)


def _oid(v: Any) -> Optional[ObjectId]:
    try:
        return ObjectId(str(v))
    except Exception:
        return None


def _ext(filename: str) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def _serialize(f) -> Dict[str, Any]:
    """Normalise a GridFSFile object OR a raw dict from `.files` into a
    JSON-safe response. Motor's `find(...).to_list()` returns dicts, while
    `open_download_stream(...)` returns GridFSFile — so we support both."""
    if isinstance(f, dict):
        _id = f.get("_id")
        filename = f.get("filename") or ""
        length = f.get("length") or 0
        upload_date = f.get("uploadDate") or f.get("upload_date")
        meta = f.get("metadata") or {}
    else:
        _id = getattr(f, "_id", None)
        filename = getattr(f, "filename", None) or ""
        length = getattr(f, "length", 0) or 0
        upload_date = getattr(f, "upload_date", None) or getattr(f, "uploadDate", None)
        meta = getattr(f, "metadata", None) or {}
    return {
        "id":            str(_id) if _id is not None else None,
        "filename":      filename,
        "length":        length,
        "upload_date":   upload_date.isoformat() if isinstance(upload_date, datetime) else upload_date,
        "content_type":  meta.get("content_type") or mimetypes.guess_type(filename)[0] or "application/octet-stream",
        "ext":           _ext(filename or ""),
        "user_email":    meta.get("user_email"),
        "notes":         meta.get("notes") or "",
        "tags":          meta.get("tags") or [],
        "ingested":      bool(meta.get("ingested_as_fixture")),
        "reinvestigated":bool(meta.get("reinvestigated")),
        "reinvestigate_history_id": meta.get("reinvestigate_history_id"),
        "sha256":        meta.get("sha256"),
    }


# ═════════════════════════════════════════════════════════════════════════════
# UPLOAD
# ═════════════════════════════════════════════════════════════════════════════
@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    """Upload a single file. Returns metadata for the stored document."""
    filename = (file.filename or "unnamed").strip()
    ext = _ext(filename)
    if ext and ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail=f"Unsupported file extension: .{ext}")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds {MAX_UPLOAD_BYTES // (1024*1024)} MB limit",
        )

    import hashlib
    sha = hashlib.sha256(data).hexdigest()

    bucket = _bucket()
    file_id = await bucket.upload_from_stream(
        filename,
        io.BytesIO(data),
        metadata={
            "user_email":    user["email"],
            "content_type":  file.content_type or mimetypes.guess_type(filename)[0]
                             or "application/octet-stream",
            "sha256":        sha,
            "uploaded_at":   datetime.now(timezone.utc).isoformat(),
            "tags":          [],
            "notes":         "",
        },
    )
    stored = await bucket.find({"_id": file_id}).to_list(1)
    return _serialize(stored[0])


# ═════════════════════════════════════════════════════════════════════════════
# LIST
# ═════════════════════════════════════════════════════════════════════════════
@router.get("/documents")
async def list_documents(
    q: str = "",
    ext: str = "",
    all: bool = False,
    limit: int = 100,
    skip: int = 0,
    user=Depends(get_current_user),
):
    """Paginated list of uploaded documents (own by default; admin can pass ?all=true)."""
    bucket = _bucket()
    query: Dict[str, Any] = {}
    is_admin = (user or {}).get("role") == "admin"
    if not (all and is_admin):
        query["metadata.user_email"] = user["email"]
    if q:
        query["filename"] = {"$regex": q, "$options": "i"}
    if ext:
        # match against extension in filename (case-insensitive)
        query["filename"] = {"$regex": rf"\.{ext.strip('.').lower()}$", "$options": "i"}

    total = await db[f"{BUCKET_NAME}.files"].count_documents(query)
    cursor = bucket.find(query).sort("uploadDate", -1).skip(max(0, skip)).limit(max(1, min(500, limit)))
    items: List[Dict[str, Any]] = []
    async for f in cursor:
        items.append(_serialize(f))
    return {"total": total, "items": items, "limit": limit, "skip": skip}


@router.get("/documents/{doc_id}")
async def get_document(doc_id: str, user=Depends(get_current_user)):
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    async for f in bucket.find({"_id": oid}):
        meta = f.metadata or {}
        if meta.get("user_email") != user["email"] and user.get("role") != "admin":
            raise HTTPException(status_code=403, detail="forbidden")
        return _serialize(f)
    raise HTTPException(status_code=404, detail="not found")


# ═════════════════════════════════════════════════════════════════════════════
# DOWNLOAD
# ═════════════════════════════════════════════════════════════════════════════
@router.get("/documents/{doc_id}/download")
async def download_document(doc_id: str, user=Depends(get_current_user)):
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    # Verify ownership
    f_meta = None
    async for f in bucket.find({"_id": oid}):
        f_meta = f
        break
    if not f_meta:
        raise HTTPException(status_code=404, detail="not found")
    meta = _gf(f_meta, 'metadata') or {}
    if meta.get("user_email") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="forbidden")

    stream = await bucket.open_download_stream(oid)
    data = await stream.read()

    ctype = meta.get("content_type") or "application/octet-stream"
    filename = _gf(f_meta, 'filename') or "download"
    return StreamingResponse(
        io.BytesIO(data),
        media_type=ctype,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/documents/{doc_id}/preview")
async def preview_document(doc_id: str, max_chars: int = 20000, user=Depends(get_current_user)):
    """Best-effort text preview. For binary formats returns a placeholder."""
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    f_meta = None
    async for f in bucket.find({"_id": oid}):
        f_meta = f
        break
    if not f_meta:
        raise HTTPException(status_code=404, detail="not found")
    meta = _gf(f_meta, 'metadata') or {}
    if meta.get("user_email") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="forbidden")

    ext = _ext(_gf(f_meta, 'filename') or "")
    stream = await bucket.open_download_stream(oid)
    data = await stream.read()

    if ext in PREVIEWABLE_TEXT_EXTS:
        try:
            text = data.decode("utf-8", errors="replace")
        except Exception:
            text = data.decode("latin-1", errors="replace")
        return {"kind": "text", "ext": ext, "content": text[:max_chars],
                "truncated": len(text) > max_chars, "length": len(text)}

    if ext == "pdf":
        try:
            from pdfminer.high_level import extract_text
            with io.BytesIO(data) as buf:
                text = extract_text(buf) or ""
            return {"kind": "text", "ext": ext, "content": text[:max_chars],
                    "truncated": len(text) > max_chars, "length": len(text)}
        except Exception:
            pass

    if ext in ("docx",):
        try:
            from docx import Document
            with io.BytesIO(data) as buf:
                doc = Document(buf)
                text = "\n".join(p.text for p in doc.paragraphs)
            return {"kind": "text", "ext": ext, "content": text[:max_chars],
                    "truncated": len(text) > max_chars, "length": len(text)}
        except Exception:
            pass

    if ext in ("xlsx",):
        try:
            from openpyxl import load_workbook
            with io.BytesIO(data) as buf:
                wb = load_workbook(buf, read_only=True, data_only=True)
                rows: List[str] = []
                for sh in wb.sheetnames[:5]:
                    ws = wb[sh]
                    rows.append(f"── Sheet: {sh} ──")
                    for r in ws.iter_rows(values_only=True, max_row=200):
                        rows.append("\t".join("" if v is None else str(v) for v in r))
            text = "\n".join(rows)
            return {"kind": "text", "ext": ext, "content": text[:max_chars],
                    "truncated": len(text) > max_chars, "length": len(text)}
        except Exception:
            pass

    return {
        "kind": "binary",
        "ext": ext,
        "content": f"Binary file ({_gf(f_meta, 'length')} bytes). Download to view.",
        "length": _gf(f_meta, 'length'),
    }


# ═════════════════════════════════════════════════════════════════════════════
# DELETE
# ═════════════════════════════════════════════════════════════════════════════
@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, user=Depends(get_current_user)):
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    async for f in bucket.find({"_id": oid}):
        meta = f.metadata or {}
        if meta.get("user_email") != user["email"] and user.get("role") != "admin":
            raise HTTPException(status_code=403, detail="forbidden")
        await bucket.delete(oid)
        return {"deleted": True, "id": doc_id}
    raise HTTPException(status_code=404, detail="not found")


# ═════════════════════════════════════════════════════════════════════════════
# INGEST AS FIXTURE (JSON only)
# ═════════════════════════════════════════════════════════════════════════════
@router.post("/documents/{doc_id}/ingest-fixture")
async def ingest_as_fixture(doc_id: str, user=Depends(get_current_user)):
    """Feed an uploaded IR-export JSON through tools/ir_export_to_fixture.py.

    On success stores the fixture path in the file metadata and returns the
    fixture body so the frontend can preview + open it in Workspace.
    """
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    f_meta = None
    async for f in bucket.find({"_id": oid}):
        f_meta = f
        break
    if not f_meta:
        raise HTTPException(status_code=404, detail="not found")
    meta = _gf(f_meta, 'metadata') or {}
    if meta.get("user_email") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="forbidden")

    ext = _ext(_gf(f_meta, 'filename') or "")
    if ext not in ("json", "jsonl"):
        raise HTTPException(status_code=400, detail="Only JSON/JSONL can be ingested as fixtures")

    stream = await bucket.open_download_stream(oid)
    data = await stream.read()
    text = data.decode("utf-8", errors="replace")

    try:
        # Write to a tmp file and invoke the converter via subprocess so we
        # reuse the exact CLI validation & schema logic without duplicating it.
        import tempfile, subprocess, sys, os
        with tempfile.NamedTemporaryFile("w", suffix=f".{ext}", delete=False, encoding="utf-8") as tf:
            tf.write(text)
            tmp_path = tf.name
        try:
            proc = subprocess.run(
                [sys.executable, "/app/backend/tools/ir_export_to_fixture.py", tmp_path,
                 "--out-dir", "/app/backend/tests/fixtures/plugin_regression"],
                capture_output=True, text=True, timeout=60,
            )
            ok = proc.returncode == 0
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    except FileNotFoundError:
        raise HTTPException(status_code=501, detail="ir_export_to_fixture.py not installed")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"ingest failed: {e}")

    if ok:
        await db[f"{BUCKET_NAME}.files"].update_one(
            {"_id": oid},
            {"$set": {"metadata.ingested_as_fixture": True,
                      "metadata.ingest_stdout": (proc.stdout or "")[:4000],
                      "metadata.ingest_ts": datetime.now(timezone.utc).isoformat()}},
        )
    return {
        "ok": ok,
        "stdout": (proc.stdout or "")[:4000],
        "stderr": (proc.stderr or "")[:2000],
        "returncode": proc.returncode,
    }


# ═════════════════════════════════════════════════════════════════════════════
# RE-INVESTIGATE (text / JSON payloads → /decode/smart)
# ═════════════════════════════════════════════════════════════════════════════
class ReinvestigateOpts(BaseModel):
    input_field: Optional[str] = None   # for JSON: extract from this key
    max_chars: int = 100000


@router.post("/documents/{doc_id}/re-investigate")
async def reinvestigate_document(
    doc_id: str,
    body: Optional[Dict[str, Any]] = Body(default=None),
    user=Depends(get_current_user),
):
    """Extract the text payload from an uploaded file and pipe it through
    the deterministic decoder pipeline (/decode/smart).

    For JSON, if `input_field` is provided it will be used as the input
    string; otherwise the whole JSON body is used as the input.
    """
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    f_meta = None
    async for f in bucket.find({"_id": oid}):
        f_meta = f
        break
    if not f_meta:
        raise HTTPException(status_code=404, detail="not found")
    meta = _gf(f_meta, 'metadata') or {}
    if meta.get("user_email") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="forbidden")

    body = body or {}
    input_field = body.get("input_field")
    max_chars = int(body.get("max_chars") or 100_000)

    stream = await bucket.open_download_stream(oid)
    data = await stream.read()

    ext = _ext(_gf(f_meta, 'filename') or "")
    # Extract text depending on the format
    text: Optional[str] = None
    if ext in PREVIEWABLE_TEXT_EXTS:
        text = data.decode("utf-8", errors="replace")
    elif ext == "pdf":
        try:
            from pdfminer.high_level import extract_text as _pdf_extract
            with io.BytesIO(data) as buf:
                text = _pdf_extract(buf) or ""
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"PDF extract failed: {e}")
    elif ext == "docx":
        try:
            from docx import Document
            with io.BytesIO(data) as buf:
                text = "\n".join(p.text for p in Document(buf).paragraphs)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"DOCX extract failed: {e}")
    elif ext == "xlsx":
        try:
            from openpyxl import load_workbook
            with io.BytesIO(data) as buf:
                wb = load_workbook(buf, read_only=True, data_only=True)
                rows: List[str] = []
                for sh in wb.sheetnames[:5]:
                    ws = wb[sh]
                    for r in ws.iter_rows(values_only=True, max_row=500):
                        rows.append("\t".join("" if v is None else str(v) for v in r))
            text = "\n".join(rows)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"XLSX extract failed: {e}")
    else:
        raise HTTPException(
            status_code=415,
            detail=f".{ext} is not supported for re-investigate; download and paste manually",
        )

    if not text or not text.strip():
        raise HTTPException(status_code=422, detail="No extractable text in file")

    # JSON with a specific input_field
    if ext in ("json", "jsonl") and input_field:
        try:
            j = json.loads(text)
            if isinstance(j, dict) and input_field in j:
                text = str(j[input_field] or "")
            elif isinstance(j, list) and j and isinstance(j[0], dict) and input_field in j[0]:
                # take first record
                text = str(j[0][input_field] or "")
        except Exception:
            pass  # fall through with raw text

    text = text[:max_chars]

    # ── ADR-004 L1 fix (owner-authorised 2026-08-10) ──────────────
    #
    # This endpoint previously called `routers.ops.decode_smart` (the
    # deep-payload decoder) with a symbol `DecodeIn` that was renamed
    # on 2026-07-20 → HTTP 500 for ~3 weeks. Worse, `decode_smart` is
    # the wrong pipeline for a DOCX incident report — it produces
    # neither Investigation Model nor Attack Story / Recommendations
    # / Executive Summary / Analyst Summary.
    #
    # L1 routes DOCX-style incident text through the EXISTING MDR
    # investigation pipeline (`v2.jobs.pipeline.run_investigation_with_progress`)
    # which already generates all of the above. NO new implementations.
    # Engine A remains authoritative. `verdict_shadow` fires so DOCX
    # cases enter the Wave 1 observation store.
    try:
        from v2.jobs.pipeline import run_investigation_with_progress
        mdr = await run_investigation_with_progress(raw=text, focus=None)
    except Exception as e:
        raise HTTPException(status_code=500,
                                    detail=f"MDR investigation pipeline failed: {e}")

    # History persistence — the MDR pipeline runs record_investigation
    # itself if the caller sets it up; otherwise we compute the hash
    # and try to find the row it recorded via the shared path.
    import hashlib
    ihash = hashlib.sha256((text or "").encode("utf-8", errors="replace")).hexdigest()
    hist = await db.investigations.find_one(
        {"user_email": user["email"], "input_hash": ihash},
        sort=[("ts", -1)],
    )
    history_id = str(hist["_id"]) if hist else None

    await db[f"{BUCKET_NAME}.files"].update_one(
        {"_id": oid},
        {"$set": {
            "metadata.reinvestigated": True,
            "metadata.reinvestigate_history_id": history_id,
            "metadata.reinvestigate_ts": datetime.now(timezone.utc).isoformat(),
            "metadata.reinvestigate_pipeline": "v2.jobs.pipeline.run_investigation_with_progress",
        }},
    )

    # ── Surface the ALREADY-GENERATED MDR outputs verbatim ────────
    # No re-computation. No new fields. Just projection of what the
    # pipeline already emits.
    return {
        "ok":                       True,
        "history_id":               history_id,
        "pipeline":                 "mdr",
        # Existing top-level outputs of run_investigation_with_progress:
        "engine":                   (mdr.get("engine") if isinstance(mdr, dict) else None),
        "final_incident_summary":   (mdr.get("final_incident_summary") if isinstance(mdr, dict) else None),
        "executive_card":           (mdr.get("executive_card") if isinstance(mdr, dict) else None),
        "investigation_model":      (mdr.get("investigation_model") if isinstance(mdr, dict) else None),
        "investigation_narrative":  (mdr.get("investigation_narrative") if isinstance(mdr, dict) else None),
        "investigation_report":     (mdr.get("investigation_report") if isinstance(mdr, dict) else None),
        # IOC / decode / OSINT pass-throughs (already computed by the pipeline):
        "iocs":                     (mdr.get("iocs") if isinstance(mdr, dict) else {}) or {},
        "mitre":                    (mdr.get("mitre") if isinstance(mdr, dict) else []) or [],
        "lolbas":                   (mdr.get("lolbas") if isinstance(mdr, dict) else []) or [],
        "decode_pipeline":          (mdr.get("decode_pipeline") if isinstance(mdr, dict) else None),
        # Phase 4 Wave 1 shadow — attaches when the pipeline produced
        # a CIO; DOCX cases now enter the observation store.
        "verdict_shadow":           (mdr.get("verdict_shadow") if isinstance(mdr, dict) else None),
    }


# ═════════════════════════════════════════════════════════════════════════════
# BATCH DECODE — extract EVERY command line from a document and run each
# through `/decode/smart`. Auto-persists each as a saved Case named
# `<file>::L<line>` so the analyst can browse them in the Case Library.
# ═════════════════════════════════════════════════════════════════════════════
_SHELL_HINT_RE = __import__("re").compile(
    r"(powershell|cmd(\.exe)?|certutil|mshta|rundll32|bitsadmin|regsvr32|"
    r"wscript|cscript|schtasks|wmic|curl\s+|wget\s+|iex\s|invoke-|"
    r"FromBase64String|System\.Convert|System\.Reflection|"
    r"http://|https://|ftp://|-e(ncodedcommand)?\s+[A-Za-z0-9+/=]{20,})",
    __import__("re").IGNORECASE,
)


def _split_command_lines(text: str, min_len: int = 24, max_lines: int = 200) -> List[str]:
    """Split extracted document text into candidate command lines.

    Rules:
      • Line-based split (\n or tab-separated cells for XLSX)
      • Keep only lines matching a shell/URL/PowerShell hint
      • Drop duplicates preserving order
      • Cap at `max_lines` to keep the batch bounded
    """
    seen, out = set(), []
    for raw in text.splitlines():
        for cell in raw.split("\t"):
            s = cell.strip()
            if len(s) < min_len or s in seen:
                continue
            if _SHELL_HINT_RE.search(s):
                seen.add(s)
                out.append(s)
                if len(out) >= max_lines:
                    return out
    return out


@router.post("/documents/{doc_id}/batch-decode")
async def batch_decode_document(
    doc_id: str,
    body: Optional[Dict[str, Any]] = Body(default=None),
    user=Depends(get_current_user),
):
    """Extract every candidate command-line from an uploaded document and
    pipe each through `/decode/smart`. Each result is auto-saved as a Case
    named `<filename>::L<index>` so the analyst can browse the whole batch
    in the Case Library and hit SIGMA / RE-DECODE per row.

    Body (all optional):
      max_lines: cap the number of extracted lines (default 200)
      min_len:   minimum candidate length in chars (default 24)
      auto_save: create workspace_cases entries (default True)
    """
    oid = _oid(doc_id)
    if not oid:
        raise HTTPException(status_code=400, detail="invalid id")
    bucket = _bucket()
    f_meta = None
    async for f in bucket.find({"_id": oid}):
        f_meta = f
        break
    if not f_meta:
        raise HTTPException(status_code=404, detail="not found")
    meta = _gf(f_meta, "metadata") or {}
    if meta.get("user_email") != user["email"] and user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="forbidden")

    body = body or {}
    max_lines = int(body.get("max_lines") or 200)
    min_len = int(body.get("min_len") or 24)
    auto_save = body.get("auto_save", True)

    ext = _ext(_gf(f_meta, "filename") or "")
    stream = await bucket.open_download_stream(oid)
    data = await stream.read()

    # Reuse the same extraction paths as re-investigate
    text: Optional[str] = None
    if ext in PREVIEWABLE_TEXT_EXTS:
        text = data.decode("utf-8", errors="replace")
    elif ext == "pdf":
        from pdfminer.high_level import extract_text as _pdf_extract
        with io.BytesIO(data) as buf:
            text = _pdf_extract(buf) or ""
    elif ext == "docx":
        from docx import Document
        with io.BytesIO(data) as buf:
            text = "\n".join(p.text for p in Document(buf).paragraphs)
    elif ext == "xlsx":
        from openpyxl import load_workbook
        with io.BytesIO(data) as buf:
            wb = load_workbook(buf, read_only=True, data_only=True)
            rows: List[str] = []
            for sh in wb.sheetnames[:5]:
                ws = wb[sh]
                for r in ws.iter_rows(values_only=True, max_row=2000):
                    rows.append("\t".join("" if v is None else str(v) for v in r))
        text = "\n".join(rows)
    else:
        raise HTTPException(status_code=415, detail=f".{ext} is not supported for batch decode")

    if not text or not text.strip():
        raise HTTPException(status_code=422, detail="No extractable text in file")

    lines = _split_command_lines(text, min_len=min_len, max_lines=max_lines)
    if not lines:
        return {"ok": False, "reason": "no candidate command lines found",
                "total_lines": len(text.splitlines())}

    from routers.ops import decode_smart
    from schemas import AutoIn as _DecodeIn

    fname = _gf(f_meta, "filename") or "doc"
    stem = fname.rsplit(".", 1)[0][:40]

    results: List[Dict[str, Any]] = []
    counts = {"malicious": 0, "partial": 0, "undecoded": 0, "benign": 0, "error": 0}

    for idx, line in enumerate(lines, start=1):
        try:
            r = await decode_smart(_DecodeIn(input=line), user=user)
            def _g(k, default=None):
                if isinstance(r, dict): return r.get(k, default)
                return getattr(r, k, default)
            vc = _g("verdict_card") or {}
            v = (vc.get("verdict") or "").lower()
            if v not in counts: v = "error"
            counts[v] = counts.get(v, 0) + 1

            # ── OUTPUT VERIFICATION (Feb 2026) ─────────────────────────
            # Don't trust the verdict alone — cross-check that the DECODED
            # content actually contains real forensic evidence.
            _out_text = str(_g("output") or "")
            _clean = _out_text
            for _tok in ("━" * 30, "▼ DECODED OUTPUT", "NIVXRAY INVESTIGATION SUMMARY"):
                _clean = _clean.split(_tok, 1)[0].strip()
            _out_bytes = _clean.encode("latin-1", errors="replace")
            _wordhits_ = _lolbas_present = _url_present = _magic_ok = False
            try:
                from ops_extended import _wordhits as _wh
                from ops_extended import _score_downstream_magic as _mag
                _wordhits_ = _wh(_out_bytes) >= 2
                _magic_ok = _mag(_out_bytes) >= 0.30
            except Exception:
                pass
            # Are LOLBAS binaries / URLs actually present in the decoded text?
            for lb in (_g("lolbas") or []):
                bn = (lb.get("binary") if isinstance(lb, dict) else lb) or ""
                if bn and bn.lower() in _clean.lower():
                    _lolbas_present = True; break
            for u in ((_g("iocs") or {}).get("urls") or []):
                if str(u).lower() in _clean.lower():
                    _url_present = True; break
            _verified = bool(_wordhits_ or _magic_ok or _lolbas_present or _url_present)
            _verify_reason = (
                "magic_bytes" if _magic_ok else
                "lolbas_in_output" if _lolbas_present else
                "url_in_output" if _url_present else
                "shell_tokens" if _wordhits_ else
                "wrapper_only_no_payload_evidence"
            )

            row = {
                "index":      idx,
                "input":      line[:200],
                "verdict":    vc.get("verdict"),
                "risk_score": vc.get("risk_score"),
                "partial":    bool(vc.get("partial")),
                "wrapper_only": bool(vc.get("wrapper_only")),
                "engine":     _g("engine"),
                "chain_len":  len(_g("layer_trace") or []),
                "urls":       (_g("iocs") or {}).get("urls") or [],
                "lolbas":     [l.get("binary") if isinstance(l, dict) else l for l in (_g("lolbas") or [])][:3],
                "mitre":      [m.get("id") for m in (_g("mitre") or [])][:6],
                # Honest verification flag — analyst can filter on this
                "verified":       _verified,
                "verify_reason":  _verify_reason,
                "output_first_120": _clean[:120],
            }
            # Auto-persist as a Case so the Case Library gets populated
            if auto_save:
                case_name = f"{stem}::L{idx:03d}"[:200]
                try:
                    from routers.cases import save_case, SaveCaseIn  # type: ignore
                    await save_case(SaveCaseIn(
                        name=case_name, input=line,
                        output=_g("output") or "",
                        engine=_g("engine") or "-",
                        confidence=_g("confidence"),
                        chain_ids=[t.get("op") if isinstance(t, dict) else t
                                    for t in (_g("layer_trace") or [])],
                        verdict=vc.get("verdict"),
                        iocs=_g("iocs") or {},
                    ), user=user)
                    row["case_name"] = case_name
                except Exception as _se:
                    row["case_save_error"] = str(_se)[:200]
            results.append(row)
        except Exception as e:
            counts["error"] += 1
            results.append({"index": idx, "input": line[:200], "error": str(e)[:200]})

    return {
        "ok": True,
        "filename": fname,
        "extracted_lines": len(lines),
        "counts": counts,
        "results": results,
    }
