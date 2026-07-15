"""HTML + DOCX exporters for the NivXRay auto-generated user guide.

Public entry-points
    generate_html(audience='user') -> str    (standalone HTML)
    generate_docx(audience='user') -> bytes  (DOCX bytes)

Both reuse `docs.generate_guide()` for the Markdown source so the three
export formats (Markdown/HTML/DOCX/PDF) never drift apart.
"""
from __future__ import annotations

import base64
import io
from datetime import datetime, timezone
from pathlib import Path as _P
from typing import Any, Dict, List

import markdown as md_lib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from docs import generate_guide, guide_stats, list_features, list_workflows

_SCREENSHOTS_DIR = _P(__file__).parent / "screenshots"


def _shots_for(doc_id: str) -> List[_P]:
    d = _SCREENSHOTS_DIR / doc_id
    if not d.exists():
        return []
    return sorted(d.glob("step_*.png"))


def _png_as_data_uri(path: _P) -> str:
    """Base64-inline a PNG so the HTML export is self-contained and works
    offline without hitting the auth-protected /api/docs/screenshots endpoint."""
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{b64}"


# ---------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------
_HTML_CSS = """
:root { color-scheme: dark; }
body {
  margin: 0; padding: 40px 60px 80px;
  font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
  background: #0f172a; color: #c9d1d9; line-height: 1.65;
  max-width: 960px; margin-left: auto; margin-right: auto;
}
header.cover {
  background: linear-gradient(135deg, #0f172a, #1e293b);
  color: white; padding: 60px 40px; margin: -40px -60px 40px;
  border-bottom: 3px solid #7ee3c9;
}
header.cover h1 { font-size: 42px; margin: 0 0 8px; letter-spacing: -0.02em; }
header.cover .sub { color: #7ee3c9; font-size: 16px; letter-spacing: 0.02em; }
header.cover .meta { color: #94a3b8; font-size: 13px; margin-top: 24px; }
h1 { color: #f8fafc; font-size: 28px; margin-top: 40px; letter-spacing: -0.01em; }
h2 { color: #7ee3c9; font-size: 20px; margin-top: 32px;
     border-bottom: 1px solid rgba(126,227,201,0.20); padding-bottom: 4px; }
h3 { color: #f8fafc; font-size: 16px; margin-top: 24px; }
p, li { font-size: 14px; }
code, pre {
  font-family: 'JetBrains Mono', 'Courier New', monospace;
  background: rgba(148,163,184,0.10); color: #f8fafc;
  padding: 2px 6px; border-radius: 3px; font-size: 13px;
}
pre { padding: 12px 14px; overflow-x: auto; border-left: 3px solid #7ee3c9; }
blockquote { border-left: 3px solid #f59e0b; padding-left: 14px;
             color: #94a3b8; font-style: italic; }
hr { border: none; border-top: 1px solid rgba(148,163,184,0.20); margin: 32px 0; }
a { color: #7ee3c9; text-decoration: none; border-bottom: 1px dotted #7ee3c9; }
footer {
  margin-top: 60px; padding-top: 20px;
  border-top: 1px solid rgba(148,163,184,0.20);
  font-size: 11px; color: #64748b; text-align: center;
}
"""


def generate_html(audience: str = "user") -> str:
    audience = audience if audience in {"user", "admin", "developer", "all"} else "user"
    guide_md = generate_guide(audience=audience)
    body = md_lib.markdown(
        guide_md,
        extensions=["fenced_code", "tables", "toc"],
        output_format="html5",
    )

    # ─── Inject screenshots after each matching <h3> heading ───────
    # Build a title → id map from the two YAML directories, then walk
    # the rendered HTML and insert `<img>` blocks (base64 data-URIs so
    # the export is self-contained).
    import re
    title_to_id: Dict[str, str] = {}
    for f in list_features():
        if f.get("title") and f.get("id"):
            title_to_id[f["title"]] = f["id"]
    for w in list_workflows():
        if w.get("title") and w.get("id"):
            title_to_id[w["title"]] = w["id"]

    def _shot_block(doc_id: str) -> str:
        shots = _shots_for(doc_id)
        if not shots:
            return ""
        imgs = []
        for i, p in enumerate(shots, 1):
            imgs.append(
                f'<figure style="margin:12px 0;padding:0;text-align:center;">'
                f'<img src="{_png_as_data_uri(p)}" '
                f'alt="{doc_id} screenshot {i}" '
                f'style="max-width:100%;height:auto;border-radius:4px;'
                f'border:1px solid rgba(148,163,184,0.20);background:#0b1220;">'
                f'<figcaption style="font-size:11px;color:#94a3b8;'
                f'margin-top:4px;">Screenshot {i}</figcaption></figure>'
            )
        return '<div class="shot-gallery">' + "".join(imgs) + "</div>"

    def _inject(match: re.Match) -> str:
        heading = match.group(0)
        title = re.sub(r"<[^>]+>", "", match.group(1)).strip()
        doc_id = title_to_id.get(title)
        if not doc_id:
            return heading
        return heading + _shot_block(doc_id)

    body = re.sub(r"<h3[^>]*>(.*?)</h3>", _inject, body, flags=re.DOTALL)

    stats = guide_stats()
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Inline the 5W1H analyst flow SVG + 3 anatomy diagrams.
    flow_svg = ""
    try:
        svg_path = _P(__file__).parent / "assets" / "analyst_flow.svg"
        if svg_path.exists():
            raw = svg_path.read_text(encoding="utf-8")
            if raw.startswith("<?xml"):
                raw = raw.split("?>", 1)[-1]
            flow_svg = (
                '<section style="margin:24px 0 32px;padding:20px;'
                'background:rgba(126,227,201,0.04);border-left:3px solid #7ee3c9;'
                'border-radius:4px;">'
                '<h2 style="margin-top:0;">Analyst Flow · 5W1H</h2>'
                '<p style="color:#94a3b8;font-size:13px;margin-top:0;">'
                'Follow the arrows. Every step answers one of the six analyst '
                'questions (What · Where · When · Why · How · Which) and loops '
                'back into the learning system.</p>'
                f'<div style="text-align:center;overflow-x:auto;">{raw}</div>'
                '</section>'
            )

        anatomy_map = [
            ("Auto-Investigate · Pipeline Anatomy", "auto_investigate_pipeline.svg"),
            ("Attack Graph · Node Anatomy",         "attack_graph_anatomy.svg"),
            ("Decoding Chain · Stage Anatomy",      "decoding_chain_anatomy.svg"),
        ]
        for heading, name in anatomy_map:
            p = _P(__file__).parent / "assets" / name
            if not p.exists():
                continue
            raw = p.read_text(encoding="utf-8")
            if raw.startswith("<?xml"):
                raw = raw.split("?>", 1)[-1]
            flow_svg += (
                '<section style="margin:24px 0 32px;padding:20px;'
                'background:rgba(167,139,250,0.04);border-left:3px solid #a78bfa;'
                'border-radius:4px;">'
                f'<h2 style="margin-top:0;">{heading}</h2>'
                f'<div style="text-align:center;overflow-x:auto;">{raw}</div>'
                '</section>'
            )
    except Exception:
        pass

    return f"""<!doctype html>
<html lang="en"><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>NivXRay — {audience.title()} Guide</title>
<style>{_HTML_CSS}</style>
</head><body>
<header class="cover">
  <h1>NivXRay</h1>
  <div class="sub">{audience.title()} Guide · CyberChef-style Decoder &amp; Threat Analysis</div>
  <div class="meta">
    {stats.get('features', 0)} features &nbsp;·&nbsp;
    {stats.get('workflows', 0)} workflows &nbsp;·&nbsp;
    {len(stats.get('categories') or [])} categories &nbsp;·&nbsp;
    Auto-generated {ts}
  </div>
</header>
{flow_svg}
<main>{body}</main>
<footer>© {datetime.now(timezone.utc).year} NivXRay · Auto-generated from
<code>docs/features/*.yaml</code> and <code>docs/workflows/*.yaml</code>.</footer>
</body></html>
"""


# ---------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------
_MINT = RGBColor(0x0F, 0x76, 0x6E)
_INK = RGBColor(0x0F, 0x17, 0x2A)
_MUTED = RGBColor(0x47, 0x55, 0x69)
_AMBER = RGBColor(0xB4, 0x53, 0x09)


def _para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
          size: int = 11, color: RGBColor = _INK, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _bullet(doc: Document, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def _feature_section(doc: Document, f: Dict[str, Any]) -> None:
    doc.add_heading(f.get("title", f.get("id", "?")), level=3)
    _para(doc,
          f"id: {f.get('id', '')}   ·   category: {f.get('category', '?')}   "
          f"·   audience: {f.get('audience', 'user')}",
          italic=True, size=9, color=_MUTED)
    if f.get("purpose"):
        _para(doc, f["purpose"], size=11)
    for label, key in (("When to use", "when_to_use"),
                       ("Supported formats", "supported_formats"),
                       ("Confidence rules", "confidence_rules"),
                       ("Common errors", "common_errors"),
                       ("Tips", "tips")):
        vals = f.get(key) or []
        if vals:
            _para(doc, label, bold=True, size=11, color=_MINT)
            _bullet(doc, [str(v) for v in vals])
    exs = f.get("examples") or []
    if exs:
        _para(doc, "Examples", bold=True, size=11, color=_MINT)
        for ex in exs:
            p = doc.add_paragraph()
            r1 = p.add_run(f"Input: {ex.get('input', '')}\n")
            r1.font.name = "Consolas"; r1.font.size = Pt(9)
            r2 = p.add_run(f"Output: {ex.get('output', '')}")
            r2.font.name = "Consolas"; r2.font.size = Pt(9)
            if ex.get("notes"):
                _para(doc, ex["notes"], italic=True, size=9, color=_AMBER)
    if f.get("related"):
        _para(doc, "Related: " + ", ".join(f["related"]),
              italic=True, size=9, color=_MUTED)
    # Inline captured screenshots
    for p in _shots_for(f.get("id", "")):
        try:
            doc.add_picture(str(p), width=Inches(6.5))
        except Exception:
            continue


def _workflow_section(doc: Document, w: Dict[str, Any]) -> None:
    doc.add_heading(w.get("title", w.get("id", "?")), level=3)
    if w.get("purpose"):
        _para(doc, w["purpose"], italic=True, color=_MUTED, size=10)
    for i, step in enumerate(w.get("steps") or [], 1):
        _para(doc, f"Step {i} — {step.get('title', '')}", bold=True, color=_INK)
        if step.get("action"):
            _para(doc, f"Action: {step['action']}", size=10)
        if step.get("expected"):
            _para(doc, f"Expected: {step['expected']}", size=10, color=_MUTED)
    # Inline captured screenshots for this workflow
    for p in _shots_for(w.get("id", "")):
        try:
            doc.add_picture(str(p), width=Inches(6.5))
        except Exception:
            continue
    if w.get("related_features"):
        _para(doc, "Related features: " + ", ".join(w["related_features"]),
              italic=True, size=9, color=_MUTED)


def generate_docx(audience: str = "user") -> bytes:
    audience = audience if audience in {"user", "admin", "developer", "all"} else "user"
    doc = Document()

    # Cover
    _para(doc, "NivXRay", bold=True, size=36, color=_INK,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"{audience.title()} Guide", size=14, color=_MINT,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "CyberChef-style decoder & threat analysis platform",
          italic=True, size=10, color=_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    stats = guide_stats()
    _para(doc,
          f"{stats.get('features', 0)} features  ·  "
          f"{stats.get('workflows', 0)} workflows  ·  "
          f"{len(stats.get('categories') or [])} categories",
          size=9, color=_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"Auto-generated {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
          size=9, color=_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Workflows
    wfs = list_workflows()
    if wfs:
        doc.add_heading("Task-Oriented Workflows", level=1)
        _para(doc,
              "Real analyst workflows — start here. Each workflow chains the "
              "features below into an investigation.",
              italic=True, color=_MUTED, size=10)
        for w in wfs:
            _workflow_section(doc, w)
        doc.add_page_break()

    # Features by category
    feats = list_features(audience=None if audience == "all" else audience)
    by_cat: Dict[str, List[Dict[str, Any]]] = {}
    for f in feats:
        by_cat.setdefault(f.get("category") or "Uncategorised", []).append(f)
    doc.add_heading("Features by Category", level=1)
    for cat in sorted(by_cat.keys()):
        doc.add_heading(cat, level=2)
        for f in by_cat[cat]:
            _feature_section(doc, f)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
