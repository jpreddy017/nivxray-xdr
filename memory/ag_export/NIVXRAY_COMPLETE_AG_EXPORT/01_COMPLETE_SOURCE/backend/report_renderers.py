"""NivXRay — Report renderers (TXT / HTML / DOCX / PDF / CSV).

Pure functions — no FastAPI, no DB. Called from routers/reports.py.
"""
from __future__ import annotations
import io
import re as _re
from html import escape as _e
from typing import Any, Dict, List

from fastapi import HTTPException
from fastapi.responses import StreamingResponse


def download(payload: bytes, filename: str, media_type: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(payload),
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Filename": filename,
            "Access-Control-Expose-Headers": "Content-Disposition, X-Filename",
        },
    )


def render_csv_report(user, ts, body, ctx) -> str:
    import csv
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["type", "id_or_kind", "value", "severity_or_tactic", "detail", "source"])
    w.writerow(["META", "generated_at", ts, "", "", "nivxray"])
    w.writerow(["META", "analyst", user["email"], "", "", "nivxray"])
    w.writerow(["META", "verdict", ctx["risk"]["verdict"], f"score={ctx['risk']['score']}", "", "heuristic"])
    fam = (ctx["description"] or {}).get("malware_family") if ctx["description"] else None
    if fam and fam.get("name"):
        w.writerow(["META", "malware_family", fam.get("name", ""), fam.get("confidence", ""), fam.get("rationale", ""), "ai"])
    for m in ctx["mitre"] or []:
        w.writerow(["MITRE", m.get("id", ""), m.get("technique", ""), m.get("tactic", ""), m.get("evidence", ""), m.get("source", "heuristic")])
    for y in ctx["yara"] or []:
        w.writerow(["YARA", y.get("rule", ""), y.get("match", ""), y.get("severity", ""), y.get("description", ""), "yara-lite"])
    for l in ctx["lolbas"] or []:
        w.writerow(["LOLBAS", l.get("binary", ""), ";".join(l.get("purposes", [])), ";".join(l.get("mitre", [])), l.get("description", ""), l.get("url", "")])
    for kind, arr in (ctx["iocs"] or {}).items():
        for v in arr or []:
            w.writerow(["IOC", kind, v, "", "", "extracted"])
    for h in ctx["ti_hits"] or []:
        w.writerow(["TI-HIT", h.get("kind", ""), h.get("value", ""), h.get("severity", ""), ";".join(h.get("tags") or []), h.get("source", "")])
    if ctx["osint"] and not ctx["osint"].get("error"):
        for ip in ctx["osint"].get("ips") or []:
            geo = ip.get("geo") or {}
            vt = ip.get("virustotal") or {}
            w.writerow(["OSINT-IP", "ip", ip["value"], geo.get("country", ""),
                        f"vt_malicious={vt.get('malicious', 0)};abuseipdb={(ip.get('abuseipdb') or {}).get('abuse_confidence_score', '')};rdns={ip.get('reverse_dns', '')}",
                        ";".join(ctx["osint"].get("sources_used") or [])])
        for d in ctx["osint"].get("domains") or []:
            vt = d.get("virustotal") or {}
            w.writerow(["OSINT-DOMAIN", "domain", d["value"], "", f"vt_malicious={vt.get('malicious', 0)};resolved={','.join(d.get('resolved_ips') or [])}", ""])
        for h in ctx["osint"].get("hashes") or []:
            vt = h.get("virustotal") or {}
            w.writerow(["OSINT-HASH", h["algorithm"], h["value"], "", f"vt_malicious={vt.get('malicious', 0)};label={vt.get('threat_label', '')}", ""])
    return buf.getvalue()


def render_docx_report(user, ts, body, ctx) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(10)

    h = doc.add_heading("NivXRay — Decoder & Threat Analysis Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {ts}\n").italic = True
    meta.add_run(f"Analyst: {user['email']}").italic = True

    doc.add_heading("Verdict", 1)
    p = doc.add_paragraph()
    p.add_run(f"{ctx['risk']['verdict']} · heuristic score {ctx['risk']['score']}/100\n").bold = True
    v = ctx["verdict"] or {}
    if v and not v.get("error"):
        p.add_run(f"AI: {v.get('verdict')} · confidence {v.get('confidence')}%\n").bold = True
        p.add_run(v.get("summary", ""))
    fam = (ctx["description"] or {}).get("malware_family") if ctx["description"] else None
    if fam and fam.get("name"):
        pf = doc.add_paragraph()
        r = pf.add_run(f"Malware family: {fam.get('name')} ({fam.get('confidence', '?')} confidence)")
        r.bold = True
        r.font.color.rgb = RGBColor(0xE2, 0x7E, 0x5D)
        doc.add_paragraph(fam.get("rationale", ""))

    doc.add_heading("Input (raw)", 1)
    doc.add_paragraph((body.input or "")[:1500])
    doc.add_heading("Decoded output", 1)
    doc.add_paragraph((body.output or "")[:3000])

    d = ctx["description"] or {}
    if d and not d.get("error"):
        doc.add_heading("AI Analysis", 1)
        if d.get("summary"):
            doc.add_paragraph(d["summary"])
        if d.get("behavior"):
            doc.add_heading("Behavior", 2)
            for b in d["behavior"]:
                doc.add_paragraph(b, style="List Bullet")
        if d.get("ioc_narrative"):
            doc.add_heading("IOC Narrative", 2)
            doc.add_paragraph(d["ioc_narrative"])
        if d.get("attribution_hints"):
            doc.add_heading("Attribution Hints", 2)
            doc.add_paragraph(d["attribution_hints"])
        if d.get("recommended_actions"):
            doc.add_heading("Recommended Actions", 2)
            for a in d["recommended_actions"]:
                doc.add_paragraph(a, style="List Bullet")

    if ctx["mitre"]:
        doc.add_heading("MITRE ATT&CK", 1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "ID"; hdr[1].text = "Technique"; hdr[2].text = "Tactic"
        hdr[3].text = "Evidence"; hdr[4].text = "Source"
        for m in ctx["mitre"]:
            row = table.add_row().cells
            row[0].text = m.get("id", ""); row[1].text = m.get("technique", "")
            row[2].text = m.get("tactic", ""); row[3].text = m.get("evidence", "")
            row[4].text = m.get("source", "heuristic")

    if ctx["lolbas"]:
        doc.add_heading("LOLBAS Matches", 1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Binary"; hdr[1].text = "Purposes"; hdr[2].text = "MITRE"; hdr[3].text = "Description"
        for l in ctx["lolbas"]:
            row = table.add_row().cells
            row[0].text = l.get("binary", "")
            row[1].text = ", ".join(l.get("purposes", []))
            row[2].text = ", ".join(l.get("mitre", []))
            row[3].text = l.get("description", "")

    if ctx["yara"]:
        doc.add_heading("YARA-lite Hits", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Rule"; hdr[1].text = "Severity"; hdr[2].text = "Description"
        for y in ctx["yara"]:
            row = table.add_row().cells
            row[0].text = y.get("rule", ""); row[1].text = y.get("severity", "")
            row[2].text = y.get("description", "")

    if ctx["iocs"] and any(v for v in ctx["iocs"].values()):
        doc.add_heading("Extracted IOCs", 1)
        for k, arr in ctx["iocs"].items():
            if not arr: continue
            doc.add_heading(k.upper(), 2)
            for v in arr:
                doc.add_paragraph(v, style="List Bullet")

    if ctx["ti_hits"]:
        doc.add_heading("Local Threat-Intel Hits", 1)
        for h in ctx["ti_hits"]:
            doc.add_paragraph(f"[{h.get('severity')}] {h.get('kind')} {h.get('value')} — source: {h.get('source')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pdf_from_html(html: str) -> bytes:
    from xhtml2pdf import pisa
    subs = {
        "var(--bg)": "#101112", "var(--sf)": "#18191b", "var(--inset)": "#0a0a0c",
        "var(--br)": "#2d3135", "var(--ac)": "#4aa890", "var(--warn)": "#e27e5d",
        "var(--hi)": "#d96c6c", "var(--tx)": "#e5e7eb", "var(--dim)": "#8b949e",
    }
    pdf_html = html
    for k, v in subs.items():
        pdf_html = pdf_html.replace(k, v)
    pdf_html = _re.sub(r":root\s*\{[^}]*\}", "", pdf_html)
    pdf_html = _re.sub(r"display\s*:\s*flex[^;]*;?", "", pdf_html)
    pdf_html = _re.sub(r"gap\s*:\s*\d+px;?", "", pdf_html)
    buf = io.BytesIO()
    result = pisa.CreatePDF(io.StringIO(pdf_html), dest=buf, encoding="utf-8")
    if result.err:
        raise HTTPException(status_code=500, detail=f"PDF render failed ({result.err} errors)")
    return buf.getvalue()


def render_text_report(user, ts, body, risk, mitre, yara, lolbas, iocs, ti_hits, osint, description, verdict):
    lines = [
        "NIVXRAY — DECODER & THREAT ANALYSIS REPORT",
        f"Generated: {ts}",
        f"Analyst:   {user['email']}",
        "=" * 68, "",
        f"VERDICT:   {risk['verdict']}   (heuristic score {risk['score']}/100)",
    ]
    if verdict and not verdict.get("error"):
        lines += [f"AI:        {verdict.get('verdict')}   ({verdict.get('confidence')}% confidence)"]
    if description and not description.get("error"):
        fam = description.get("malware_family") or {}
        if fam.get("name"):
            lines += [f"FAMILY:    {fam.get('name')}  ({fam.get('confidence','?')} confidence)"]
    lines += ["", "INPUT (first 400 chars):", (body.input or "")[:400], "",
              "DECODED OUTPUT (first 1500 chars):", (body.output or "")[:1500], ""]
    if description and not description.get("error"):
        lines += ["── AI EXECUTIVE SUMMARY ──", description.get("summary", ""), ""]
        if description.get("behavior"):
            lines += ["── BEHAVIOR ──"] + [f"  · {b}" for b in description["behavior"]] + [""]
        if description.get("ioc_narrative"):
            lines += ["── IOC NARRATIVE ──", description["ioc_narrative"], ""]
        if description.get("attribution_hints"):
            lines += ["── ATTRIBUTION HINTS ──", description["attribution_hints"], ""]
        if description.get("recommended_actions"):
            lines += ["── RECOMMENDED ACTIONS ──"] + [f"  · {a}" for a in description["recommended_actions"]] + [""]
    lines += ["── MITRE ATT&CK ──"]
    for m in mitre or []:
        line = f"  - {m['id']}  {m['technique']}   [{m['tactic']}]"
        if m.get("evidence"): line += f"\n     evidence: {m['evidence']}"
        lines.append(line)
    if not mitre: lines.append("  (none)")
    lines += ["", "── LOLBAS MATCHES ──"]
    for l in lolbas or []:
        lines += [f"  · {l['binary']}   purposes={','.join(l['purposes'])}   mitre={','.join(l['mitre'])}",
                  f"     {l['description']}",
                  f"     snippet: {l['snippet']}"]
    if not lolbas: lines.append("  (none)")
    lines += ["", "── YARA-LITE HITS ──"]
    for y in yara or []:
        lines.append(f"  - [{y['severity'].upper()}] {y['rule']}: {y['description']}")
    if not yara: lines.append("  (none)")
    lines += ["", "── IOCs ──"]
    for k, v in (iocs or {}).items():
        if v:
            lines.append(f"  {k}:")
            for item in v: lines.append(f"    - {item}")
    lines += ["", "── LOCAL THREAT-INTEL HITS ──"]
    for h in ti_hits or []:
        lines.append(f"  - [{h['severity']}] {h['kind']} {h['value']}  (source: {h['source']})")
    if not ti_hits: lines.append("  (none)")
    if osint:
        lines += ["", "── OSINT ENRICHMENT ──", f"  sources: {', '.join(osint.get('sources_used') or [])}"]
        for ip in osint.get("ips", []) or []:
            geo = ip.get("geo") or {}
            lines.append(f"  IP {ip['value']}: {geo.get('country','?')} / {geo.get('isp','?')}"
                         + (f"  rDNS={ip.get('reverse_dns')}" if ip.get('reverse_dns') else ""))
            if ip.get("virustotal"):
                lines.append(f"     VT: {ip['virustotal'].get('malicious',0)} malicious")
            if ip.get("abuseipdb"):
                lines.append(f"     AbuseIPDB: {ip['abuseipdb'].get('abuse_confidence_score',0)}% confidence")
    lines += ["", "=" * 68, "End of report."]
    return "\n".join(lines)


def render_html_report(user, ts, body, risk, mitre, yara, lolbas, iocs, ti_hits, osint, description, verdict):
    def block(title, body_html):
        return f'<section><h2>{_e(title)}</h2><div class="card">{body_html}</div></section>'
    fam = (description or {}).get("malware_family") or {}
    parts = []
    parts.append(f'''<!doctype html><html><head><meta charset="utf-8"><title>NivXRay report</title>
<style>
:root {{ --bg:#101112; --sf:#18191b; --inset:#0a0a0c; --br:#2d3135; --ac:#4aa890; --warn:#e27e5d; --hi:#d96c6c; --tx:#e5e7eb; --dim:#8b949e; }}
* {{ box-sizing: border-box; }}
body {{ background:var(--bg); color:var(--tx); font-family: Chivo, ui-sans-serif, sans-serif; margin:0; padding:32px; line-height:1.55; }}
.mono {{ font-family: 'JetBrains Mono', ui-monospace, monospace; }}
.hdr {{ display:flex; align-items:center; gap:14px; padding-bottom:16px; border-bottom:1px solid var(--br); }}
.hdr .logo {{ width:24px; height:24px; border:1px solid var(--ac); position:relative; }}
.hdr .logo::before {{ content:''; position:absolute; inset:6px; background:var(--ac); }}
.hdr h1 {{ font-weight:900; letter-spacing:0.14em; margin:0; }}
.hdr h1 span {{ color:var(--ac); }}
.meta {{ color:var(--dim); font-size:12px; margin-left:auto; text-align:right; }}
section {{ margin-top:26px; }}
section h2 {{ color:var(--ac); font-size:11px; letter-spacing:0.22em; margin:0 0 10px 0; }}
.card {{ background:var(--sf); border:1px solid var(--br); padding:16px; }}
.badge {{ display:inline-block; padding:3px 7px; border:1px solid var(--br); font-family:'JetBrains Mono',monospace; font-size:10px; letter-spacing:0.06em; margin-right:6px; }}
.badge.hi {{ color:var(--hi); border-color:var(--hi); background:rgba(217,108,108,0.1); }}
.badge.med {{ color:var(--warn); border-color:var(--warn); background:rgba(226,126,93,0.1); }}
.badge.low {{ color:#c0ca33; border-color:#c0ca33; background:rgba(192,202,51,0.1); }}
.badge.safe {{ color:var(--ac); border-color:var(--ac); background:rgba(74,168,144,0.1); }}
.badge.critical {{ color:var(--hi); border-color:var(--hi); background:rgba(217,108,108,0.2); }}
.badge.high {{ color:var(--hi); border-color:var(--hi); }}
.badge.medium {{ color:var(--warn); border-color:var(--warn); }}
pre {{ background:var(--inset); padding:12px; border:1px solid var(--br); overflow-x:auto; font-size:11px; color:var(--tx); white-space:pre-wrap; word-break:break-all; margin:0; }}
table {{ width:100%; border-collapse:collapse; font-family:'JetBrains Mono',monospace; font-size:11px; }}
th,td {{ text-align:left; padding:6px 8px; border-bottom:1px solid var(--br); vertical-align:top; }}
th {{ color:var(--dim); font-weight:700; letter-spacing:0.12em; font-size:10px; }}
ul {{ margin:0; padding-left:20px; }}
ul li {{ margin:4px 0; }}
.verdict {{ display:flex; align-items:center; gap:14px; }}
.verdict .score {{ font-size:40px; font-weight:900; color:var(--ac); }}
.warn {{ color:var(--warn); }}
.hi {{ color:var(--hi); }}
</style></head><body>
<div class="hdr">
  <div class="logo"></div>
  <h1>NIVX<span>RAY</span> · DECODER & THREAT ANALYSIS REPORT</h1>
  <div class="meta">Generated: {_e(ts)}<br>Analyst: {_e(user['email'])}</div>
</div>''')
    vv = verdict or {}
    fam_line = ""
    if fam.get("name"):
        fam_line = f"<div class='mono' style='margin-top:6px;'>Family: <span class='warn'>{_e(fam.get('name'))}</span> <span class='badge {fam.get('confidence','low')}'>{_e(fam.get('confidence','?'))} confidence</span></div>"
    parts.append(block("VERDICT", f"""
<div class='verdict'>
  <div class='score'>{_e(str(risk['score']))}<span style='font-size:16px; color:var(--dim);'>/100</span></div>
  <div>
    <div><span class='badge {_e(risk['level'])}'>{_e(risk['verdict'])}</span>
    {"<span class='badge hi'>AI: " + _e(vv.get('verdict','?')) + " " + _e(str(vv.get('confidence','?'))) + "%</span>" if vv and not vv.get('error') else ''}</div>
    {fam_line}
    {"<div class='mono' style='margin-top:6px; color:var(--dim);'>" + _e(vv.get('summary','')) + "</div>" if vv and not vv.get('error') else ''}
  </div>
</div>"""))

    parts.append(block("INPUT", f"<pre class='mono'>{_e((body.input or '')[:2000])}</pre>"))
    parts.append(block("DECODED OUTPUT", f"<pre class='mono'>{_e((body.output or '')[:4000])}</pre>"))

    if description and not description.get("error"):
        d = description
        blocks = []
        if d.get("summary"): blocks.append(f"<p class='mono'>{_e(d['summary'])}</p>")
        if d.get("behavior"): blocks.append("<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>BEHAVIOR</h3><ul class='mono'>" + "".join(f"<li>{_e(b)}</li>" for b in d['behavior']) + "</ul>")
        if d.get("ioc_narrative"): blocks.append(f"<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>IOC NARRATIVE</h3><p class='mono' style='color:var(--dim);'>{_e(d['ioc_narrative'])}</p>")
        if d.get("attribution_hints"): blocks.append(f"<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>ATTRIBUTION HINTS</h3><p class='mono' style='color:var(--dim);'>{_e(d['attribution_hints'])}</p>")
        if d.get("recommended_actions"): blocks.append("<h3 class='mono' style='color:var(--warn);font-size:10px;letter-spacing:0.18em;'>RECOMMENDED ACTIONS</h3><ul class='mono' style='color:var(--ac);'>" + "".join(f"<li>{_e(a)}</li>" for a in d['recommended_actions']) + "</ul>")
        parts.append(block("AI ANALYSIS", "".join(blocks)))

    if mitre:
        rows = "".join(
            f"<tr><td><a href='https://attack.mitre.org/techniques/{m['id'].replace('.','/')}/' target='_blank' style='color:var(--ac);'>{_e(m['id'])}</a></td><td>{_e(m.get('technique',''))}</td><td>{_e(m.get('tactic',''))}</td><td>{_e(m.get('evidence',''))}</td><td><span class='badge'>{_e(m.get('source','heuristic'))}</span></td></tr>"
            for m in mitre
        )
        parts.append(block("MITRE ATT&CK", f"<table><tr><th>ID</th><th>Technique</th><th>Tactic</th><th>Evidence</th><th>Source</th></tr>{rows}</table>"))

    if lolbas:
        rows = "".join(
            f"<tr><td class='warn'>{_e(l['binary'])}</td><td>{', '.join(_e(p) for p in l['purposes'])}</td><td>{', '.join(_e(t) for t in l['mitre'])}</td><td>{_e(l['description'])}</td><td><a style='color:var(--ac);' href='{_e(l['url'])}' target='_blank'>docs</a></td></tr>"
            for l in lolbas
        )
        parts.append(block("LOLBAS", f"<table><tr><th>Binary</th><th>Purposes</th><th>MITRE</th><th>Description</th><th></th></tr>{rows}</table>"))

    if yara:
        rows = "".join(
            f"<tr><td>{_e(y['rule'])}</td><td><span class='badge {_e(y['severity'])}'>{_e(y['severity'])}</span></td><td>{_e(y['description'])}</td><td class='mono' style='color:var(--dim);'>{_e(y['match'][:80])}</td></tr>"
            for y in yara
        )
        parts.append(block("YARA-LITE HITS", f"<table><tr><th>Rule</th><th>Severity</th><th>Description</th><th>Match</th></tr>{rows}</table>"))

    ioc_rows = []
    for k, v in (iocs or {}).items():
        for item in v or []:
            ioc_rows.append(f"<tr><td>{_e(k)}</td><td class='mono'>{_e(item)}</td></tr>")
    if ioc_rows:
        parts.append(block("EXTRACTED IOCs", f"<table><tr><th>Kind</th><th>Value</th></tr>{''.join(ioc_rows)}</table>"))

    if ti_hits:
        rows = "".join(
            f"<tr><td>{_e(h.get('kind',''))}</td><td class='mono'>{_e(h.get('value',''))}</td><td><span class='badge {_e(h.get('severity','low'))}'>{_e(h.get('severity',''))}</span></td><td>{_e(h.get('source',''))}</td></tr>"
            for h in ti_hits
        )
        parts.append(block("LOCAL THREAT-INTEL HITS", f"<table><tr><th>Kind</th><th>Value</th><th>Severity</th><th>Source</th></tr>{rows}</table>"))

    if osint and not osint.get("error"):
        html_bits = [f"<div class='mono' style='color:var(--dim);margin-bottom:8px;'>Sources: {_e(', '.join(osint.get('sources_used') or []))}</div>"]
        for ip in osint.get("ips") or []:
            geo = ip.get("geo") or {}
            vt = ip.get("virustotal") or {}
            ab = ip.get("abuseipdb") or {}
            html_bits.append(f"<div class='card' style='margin-bottom:8px;'><b class='mono' style='color:var(--ac);'>{_e(ip['value'])}</b> — {_e(geo.get('country',''))} · {_e(geo.get('isp',''))}"
                             + (f" · rDNS={_e(ip.get('reverse_dns',''))}" if ip.get('reverse_dns') else "")
                             + (f"<br>VT: <span class='hi'>{vt.get('malicious',0)} malicious</span>, {vt.get('suspicious',0)} suspicious" if vt else "")
                             + (f"<br>AbuseIPDB: {ab.get('abuse_confidence_score',0)}% confidence, {ab.get('total_reports',0)} reports" if ab else "")
                             + "</div>")
        parts.append(block("OSINT ENRICHMENT", "".join(html_bits)))

    parts.append('</body></html>')
    return "".join(parts)
