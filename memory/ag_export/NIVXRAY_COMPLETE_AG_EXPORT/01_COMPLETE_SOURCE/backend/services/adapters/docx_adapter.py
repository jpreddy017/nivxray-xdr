"""DOCX Evidence Adapter — Phase 3A.

Extracts every forensically valuable component of a Word document
(DOCX / DOCM) and emits them as canonical IEP artifacts.  Rule R8:
this adapter never reasons about *what the document means* — it only
reports what is there.

Scope (per frozen architecture doc):

  · Paragraphs, tables, headers/footers
  · Hyperlinks
  · Comments, tracked changes
  · Document properties (core + custom)
  · External template references
  · Embedded OLE objects
  · .docm macros (VBA)
  · Embedded packages (arbitrary bin attachments inside the ZIP part)

Every artifact carries a ``docx.<part>[.N]`` source_ref so R6 is met.
"""
from __future__ import annotations

import io
import re
import zipfile
from typing import Any, Dict, List, Optional
from xml.etree import ElementTree as ET

from models.iep import (
    IEPArtifact,
    IEPContent,
    IEPRelationship,
    IEPSource,
    IEPWarning,
    RelationshipType,
)
from services.ida.artifact_splitter import split_artifacts

from .base import EvidenceAdapter


_NS = {
    "w":  "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r":  "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
}


class DOCXAdapter(EvidenceAdapter):
    name         = "adapter.docx"
    version      = "1.0"
    capabilities = [
        "paragraphs", "tables", "headers_footers", "hyperlinks",
        "comments", "tracked_changes", "document_properties",
        "custom_properties", "external_templates", "embedded_ole",
        "macros_vba", "embedded_packages",
    ]

    # DOCX / DOCM files are ZIPs whose first bytes are `PK\x03\x04`.
    # We disambiguate from raw ZIP by inspecting the archive contents.
    _MAGIC = b"PK\x03\x04"

    # ── Detection ────────────────────────────────────────────────────
    def can_handle(self, raw: Any) -> bool:
        if not isinstance(raw, (bytes, bytearray)):
            return False
        if raw[:4] != self._MAGIC:
            return False
        # Quick peek: an Office DOCX must contain `word/document.xml`.
        try:
            with zipfile.ZipFile(io.BytesIO(bytes(raw))) as z:
                names = z.namelist()
                return "word/document.xml" in names
        except Exception:
            return False

    # ── Extraction ───────────────────────────────────────────────────
    def extract(self, raw: Any) -> IEPContent:
        data = bytes(raw)
        blocks: List[Dict[str, Any]] = []
        text_parts: List[str] = []
        info: Dict[str, Any] = {
            "hyperlinks":         [],
            "comments":           [],
            "tracked_changes":    [],
            "document_props":     {},
            "custom_props":       {},
            "external_templates": [],
            "embedded_ole":       [],
            "macros":             [],
            "embedded_packages":  [],
            "warnings":           [],
        }

        # 1. python-docx for paragraphs / tables / hyperlinks / comments
        try:
            import docx  # python-docx
            doc = docx.Document(io.BytesIO(data))
            # Paragraphs
            for i, p in enumerate(doc.paragraphs, start=1):
                t = (p.text or "").strip()
                if t:
                    blocks.append({"kind": "paragraph", "index": i, "text": t})
                    text_parts.append(t)
            # Tables (flatten cell text into a joined string per row)
            for ti, table in enumerate(doc.tables, start=1):
                rows: List[List[str]] = []
                for row in table.rows:
                    rows.append([(c.text or "").strip() for c in row.cells])
                    text_parts.append(" | ".join(rows[-1]))
                blocks.append({"kind": "table", "index": ti, "rows": rows})
            # Headers / footers
            for si, section in enumerate(doc.sections, start=1):
                for part_name, part in (("header", section.header),
                                              ("footer", section.footer)):
                    text_hf = "\n".join(p.text for p in part.paragraphs
                                                 if (p.text or "").strip())
                    if text_hf:
                        blocks.append({"kind": part_name, "section": si,
                                          "text": text_hf})
                        text_parts.append(text_hf)
            # Document core properties
            cp = doc.core_properties
            info["document_props"] = {
                "author":       cp.author,
                "title":        cp.title,
                "subject":      cp.subject,
                "keywords":     cp.keywords,
                "last_modified_by": cp.last_modified_by,
                "revision":     cp.revision,
                "created":      cp.created.isoformat() if cp.created else None,
                "modified":     cp.modified.isoformat() if cp.modified else None,
                "category":     cp.category,
                "comments":     cp.comments,
                "content_status": cp.content_status,
            }
        except Exception as e:
            info["warnings"].append({
                "severity": "warn", "code": "docx_parse_failed",
                "message": f"python-docx failed: {e}",
            })

        # 2. Direct ZIP inspection for hyperlinks / OLE / macros /
        #    comments / tracked changes / custom props / template refs /
        #    embedded packages.  python-docx doesn't expose these.
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                names = z.namelist()
                # Hyperlinks (from relationships)
                if "word/_rels/document.xml.rels" in names:
                    self._parse_rels(z, "word/_rels/document.xml.rels", info)
                # Comments
                if "word/comments.xml" in names:
                    self._parse_comments(z, info)
                # Tracked changes (ins/del elements in document.xml)
                if "word/document.xml" in names:
                    self._parse_tracked_changes(z, info)
                # Custom document properties
                if "docProps/custom.xml" in names:
                    self._parse_custom_props(z, info)
                # External template reference (settings.xml)
                if "word/settings.xml" in names:
                    self._parse_settings(z, info)
                # Macros / OLE / embedded packages
                self._scan_embeddings(z, info)
        except Exception as e:
            info["warnings"].append({
                "severity": "warn", "code": "docx_zip_scan_failed",
                "message":  f"DOCX ZIP scan failed: {e}",
            })

        content = IEPContent(text="\n".join(text_parts), blocks=blocks)
        content.__dict__["_docx"] = info
        return content

    # ── Helpers ──────────────────────────────────────────────────────
    def _parse_rels(self, z, path, info):
        try:
            with z.open(path) as f:
                tree = ET.parse(f)
            for rel in tree.getroot():
                rtype = (rel.attrib.get("Type") or "").lower()
                target = rel.attrib.get("Target") or ""
                mode = (rel.attrib.get("TargetMode") or "").lower()
                if not target:
                    continue
                if "hyperlink" in rtype:
                    info["hyperlinks"].append({"target": target, "mode": mode})
                elif "attachedtemplate" in rtype or "template" in rtype:
                    info["external_templates"].append(target)
                elif "oleobject" in rtype:
                    info["embedded_ole"].append(target)
                elif "package" in rtype:
                    info["embedded_packages"].append(target)
        except Exception as e:
            info["warnings"].append({
                "severity": "warn", "code": "docx_rels_parse_failed",
                "message":  f"rels parse failed: {e}",
            })

    def _parse_comments(self, z, info):
        try:
            with z.open("word/comments.xml") as f:
                tree = ET.parse(f)
            for c in tree.iter(f"{{{_NS['w']}}}comment"):
                author = c.attrib.get(f"{{{_NS['w']}}}author") or ""
                date   = c.attrib.get(f"{{{_NS['w']}}}date")   or ""
                text = "".join(t.text or "" for t in
                                 c.iter(f"{{{_NS['w']}}}t"))
                info["comments"].append({
                    "author": author, "date": date, "text": text.strip(),
                })
        except Exception as e:
            info["warnings"].append({
                "severity": "warn", "code": "docx_comments_parse_failed",
                "message":  f"comments parse failed: {e}",
            })

    def _parse_tracked_changes(self, z, info):
        try:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
            for tag in ("ins", "del"):
                for e in tree.iter(f"{{{_NS['w']}}}{tag}"):
                    author = e.attrib.get(f"{{{_NS['w']}}}author") or ""
                    date   = e.attrib.get(f"{{{_NS['w']}}}date")   or ""
                    txt = "".join(t.text or "" for t in
                                    e.iter(f"{{{_NS['w']}}}t"))
                    info["tracked_changes"].append({
                        "type":   tag,
                        "author": author,
                        "date":   date,
                        "text":   txt.strip(),
                    })
        except Exception:
            pass

    def _parse_custom_props(self, z, info):
        try:
            with z.open("docProps/custom.xml") as f:
                tree = ET.parse(f)
            for p in tree.getroot():
                name = p.attrib.get("name") or ""
                val_e = list(p)
                val = val_e[0].text if val_e else None
                if name:
                    info["custom_props"][name] = val
        except Exception:
            pass

    def _parse_settings(self, z, info):
        try:
            with z.open("word/settings.xml") as f:
                data = f.read().decode("utf-8", errors="ignore")
            # attachedTemplate w:val="path"
            for m in re.finditer(
                r'attachedTemplate[^>]*w:val="([^"]+)"', data,
            ):
                info["external_templates"].append(m.group(1))
        except Exception:
            pass

    def _scan_embeddings(self, z, info):
        for name in z.namelist():
            lname = name.lower()
            if lname.startswith("word/vbaproject") or lname.endswith(".bin") and "vba" in lname:
                info["macros"].append({"part": name})
            elif lname.endswith(".bin") and "embeddings" in lname:
                info["embedded_ole"].append(name)
            elif "embeddings/" in lname and not lname.endswith("/"):
                info["embedded_packages"].append(name)

    # ── Normalization ────────────────────────────────────────────────
    def normalize(self, content: IEPContent) -> List[IEPArtifact]:
        info = getattr(content, "_docx", {}) or {}
        out: List[IEPArtifact] = []

        # Body text through the deterministic splitter
        for b in content.blocks or []:
            src = f"docx.{b.get('kind')}.{b.get('index', b.get('section', ''))}".rstrip(".")
            body = b.get("text") or (" | ".join(" ".join(r) for r in b.get("rows") or []))
            for a in (split_artifacts(body) or []):
                t = self._map_type(getattr(a, "type", None))
                v = getattr(a, "value", None)
                if not (t and v):
                    continue
                out.append(IEPArtifact(
                    type=t, value=v,
                    canonical=getattr(a, "canonical", None) or None,
                    confidence=getattr(a, "confidence", 1.0) or 1.0,
                    source_ref=src,
                ))

        # Hyperlinks
        for h in info.get("hyperlinks") or []:
            tgt = h.get("target") or ""
            if tgt.startswith(("http://", "https://")):
                out.append(IEPArtifact(
                    type="url", value=tgt,
                    source_ref="docx.hyperlink",
                    attributes={"origin": "docx_hyperlink"},
                ))

        # External templates
        for tpl in info.get("external_templates") or []:
            out.append(IEPArtifact(
                type=("url" if tpl.startswith(("http://", "https://")) else "file_path"),
                value=tpl,
                source_ref="docx.external_template",
                tags=["docx_external_template"],
            ))

        # Embedded OLE, packages, macros
        for kind, key, tag in (
            ("embedded_ole",      "file_path", "docx_ole"),
            ("embedded_packages", "file_path", "docx_embedded_package"),
            ("macros",            "file_path", "docx_macro"),
        ):
            for e in info.get(kind) or []:
                path = e if isinstance(e, str) else e.get("part") or ""
                if not path:
                    continue
                out.append(IEPArtifact(
                    type="file_path", value=path,
                    source_ref=f"docx.{kind}",
                    tags=[tag],
                ))

        # Comments — surface as unknown-type artifacts with the text
        # (Rule R8: no inference, just evidence)
        for i, c in enumerate(info.get("comments") or [], start=1):
            body = c.get("text") or ""
            if body:
                out.append(IEPArtifact(
                    type="unknown", value=body[:400],
                    source_ref=f"docx.comment.{i}",
                    tags=["docx_comment"],
                    attributes={"author": c.get("author"), "date": c.get("date")},
                ))
        # Tracked changes
        for i, tc in enumerate(info.get("tracked_changes") or [], start=1):
            body = tc.get("text") or ""
            if body:
                out.append(IEPArtifact(
                    type="unknown", value=body[:400],
                    source_ref=f"docx.tracked.{i}",
                    tags=[f"docx_tracked_{tc.get('type')}"],
                    attributes={"author": tc.get("author"), "date": tc.get("date")},
                ))
        return out

    # ── Relationships (R8) ───────────────────────────────────────────
    def discover_relationships(self, content, artifacts):
        info = getattr(content, "_docx", {}) or {}
        rels: List[IEPRelationship] = []
        doc = "docx.document"
        for h in info.get("hyperlinks") or []:
            tgt = h.get("target") or ""
            if not tgt:
                continue
            rels.append(IEPRelationship(
                from_ref=doc, to_ref=tgt,
                verb=RelationshipType.CONTAINS,
                source_ref="docx.hyperlink",
            ))
        for tpl in info.get("external_templates") or []:
            rels.append(IEPRelationship(
                from_ref=doc, to_ref=tpl,
                verb=RelationshipType.REFERENCES,
                source_ref="docx.external_template",
            ))
        for p in info.get("embedded_ole") or []:
            rels.append(IEPRelationship(
                from_ref=doc, to_ref=(p if isinstance(p, str) else p.get("part", "")),
                verb=RelationshipType.EMBEDS,
                source_ref="docx.embedded_ole",
            ))
        for p in info.get("embedded_packages") or []:
            rels.append(IEPRelationship(
                from_ref=doc, to_ref=p,
                verb=RelationshipType.ATTACHES,
                source_ref="docx.embedded_package",
            ))
        for m in info.get("macros") or []:
            rels.append(IEPRelationship(
                from_ref=doc, to_ref=m.get("part") if isinstance(m, dict) else m,
                verb=RelationshipType.EMBEDS,
                source_ref="docx.macro",
            ))
        return rels

    # ── Adapter-level warnings ───────────────────────────────────────
    def validate(self, iep) -> List[IEPWarning]:
        info = getattr(iep.content, "_docx", {}) or {}
        out: List[IEPWarning] = [IEPWarning(**w)
                                    for w in info.get("warnings") or []]
        if info.get("macros"):
            out.append(IEPWarning(
                severity="warn", code="docx_contains_macros",
                message=f"DOCX contains {len(info['macros'])} macro part(s).",
            ))
        if info.get("embedded_ole"):
            out.append(IEPWarning(
                severity="warn", code="docx_contains_ole",
                message=f"DOCX contains {len(info['embedded_ole'])} OLE object(s).",
            ))
        if info.get("external_templates"):
            out.append(IEPWarning(
                severity="warn", code="docx_external_template",
                message=(f"DOCX references {len(info['external_templates'])} "
                         "external template(s)."),
            ))
        if info.get("embedded_packages"):
            out.append(IEPWarning(
                severity="info", code="docx_embedded_package",
                message=f"DOCX carries {len(info['embedded_packages'])} embedded package(s).",
            ))
        return out

    # ── Recursion — embedded assets each become a child IEP ──────────
    def recurse(self, iep) -> List[IEPArtifact]:
        return [a for a in iep.artifacts
                if any(t in (a.tags or []) for t in (
                    "docx_ole", "docx_embedded_package", "docx_macro",
                ))]

    # ── Source detection ─────────────────────────────────────────────
    def _infer_source(self, raw: Any) -> IEPSource:
        import hashlib
        data = bytes(raw)
        return IEPSource(
            kind="docx",
            size_bytes=len(data),
            sha256=hashlib.sha256(data).hexdigest(),
            mime_type=("application/vnd.openxmlformats-officedocument"
                       ".wordprocessingml.document"),
        )

    # ── Splitter type → canonical IEP artifact type ──────────────────
    _TYPE_MAP = {
        "command":       "command",
        "url":           "url",
        "ip":            "ip",
        "domain":        "domain",
        "hash":          "hash",
        "file_path":     "file_path",
        "registry_key":  "registry_key",
        "email":         "email_address",
        "cve":           "cve",
    }
    def _map_type(self, t: Any) -> str:
        if not t:
            return "unknown"
        return self._TYPE_MAP.get(t, t)

    # ── Override make_iep to merge document properties into metadata ─
    def make_iep(self, raw, *, source=None, parent_iep_id=None,
                   pipeline_depth=0, metadata=None):
        content       = self.extract(raw)
        artifacts     = self.normalize(content)
        relationships = self.discover_relationships(content, artifacts)
        src           = source or self._infer_source(raw)
        info          = getattr(content, "_docx", {}) or {}
        md = dict(metadata or {})
        if info.get("document_props"):
            md["docx"] = {"document_props": info["document_props"],
                            "custom_props":  info.get("custom_props") or {}}
        md.setdefault("adapter", {})
        md["adapter"].update({
            "id":           f"{self.name}@{self.version}",
            "name":         self.name,
            "version":      self.version,
            "capabilities": list(self.capabilities),
        })
        from models.iep import make_iep as _mk
        iep = _mk(
            source=src, content=content, artifacts=artifacts,
            relationships=relationships, metadata=md,
            adapter=self.name, adapter_version=self.version,
            parent_iep_id=parent_iep_id, pipeline_depth=pipeline_depth,
        )
        iep.warnings.extend(self.validate(iep))
        iep.metadata.data["adapter"]["warnings"] = [w.code for w in iep.warnings]
        return iep
